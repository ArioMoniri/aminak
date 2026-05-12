"""Step 7 — 3Dmol viewers per model + summary.json + DOCX report.

Builds:
- 10_modeller/07_viewers/model{i}.html (model + crystal overlay, 3Dmol)
- viewers/modeller_model{i}.html (mirror for GitHub Pages)
- viewers/index.html updated with new "Homology models" section
- 10_modeller/summary.json
- 10_modeller/README_PHASE6.md
- 09e_report_v5/report_PHASE6.docx
"""
from __future__ import annotations

import csv
import json
import shutil
import sys
from datetime import datetime
from html import escape
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from _common import (  # noqa: E402
    MODELS_DIR, PHASE_DIR, PROJECT_DIR, STEP1_DIR, STEP2_DIR,
    STEP4_DIR, STEP5_DIR, STEP6_DIR, STEP7_DIR, VIEWERS_DIR, png_ok,
    setup_logger,
)

LOG = setup_logger("step7_viewers_report")

VIEWER_TEMPLATE = """<!DOCTYPE html>
<html lang="en"><head>
<meta charset="utf-8">
<title>{title}</title>
<script src="https://3Dmol.csb.pitt.edu/build/3Dmol-min.js"></script>
<style>
 body {{ font-family: system-ui, -apple-system, Segoe UI, sans-serif;
         margin: 0; padding: 0 16px 12px 16px; background:#0c1116; color:#e8eef5; }}
 h1 {{ font-size: 17px; font-weight: 600; margin: 12px 0 4px; color:#9ad5ff; }}
 .meta {{ font-size: 12px; color:#aab8c8; margin: 0 0 8px; line-height:1.4; }}
 .panel {{ background:#1a2230; border:1px solid #283344; border-radius:8px;
          padding:0; overflow:hidden; }}
 .legend {{ font-size: 11.5px; color:#cad5e6; padding: 6px 12px; border-top:1px solid #283344; }}
 .legend span {{ display:inline-block; padding: 2px 8px; border-radius: 4px;
                margin: 2px 4px 2px 0; }}
 .nav a {{ color:#9ad5ff; margin-right: 12px; font-size: 12px; text-decoration:none; }}
 .nav a:hover {{ text-decoration:underline; }}
 .scores {{ font-size: 11.5px; color:#aab8c8; margin: 6px 0; }}
 .scores code {{ color:#cfe6ff; }}
</style>
</head><body>
<div class="nav" style="margin: 10px 0 4px;">
  <a href="index.html">← Index</a>
  <a href="https://github.com/ArioMoniri/aminak">GitHub</a>
</div>
<h1>{title}</h1>
<div class="meta">{description}</div>
<div class="scores">{scores_html}</div>
<div class="panel">
 <div id="viewer" style="width:100%; max-width:800px; height:560px; margin:auto; position:relative;"></div>
 <div class="legend">
  <span style="background:#3b1d5a;color:#dccfff;">1HVY chain A (crystal, magenta)</span>
  <span style="background:#1d5a3b;color:#cfffd6;">Homology model (green)</span>
 </div>
</div>
<script>
const crystalPdb = `{crystal_pdb}`;
const modelPdb = `{model_pdb}`;
const viewer = $3Dmol.createViewer("viewer", {{ backgroundColor: "#0c1116" }});
viewer.addModel(crystalPdb, "pdb");
viewer.setStyle({{ model: 0 }}, {{ cartoon: {{ color: "magenta" }} }});
viewer.addModel(modelPdb, "pdb");
viewer.setStyle({{ model: 1 }}, {{ cartoon: {{ color: "green" }} }});
viewer.zoomTo();
viewer.render();
</script>
</body></html>
"""


def build_viewer_html(model_pdb_path: Path, crystal_pdb_path: Path, model_id: int,
                      scores: dict) -> str:
    title = f"Homology model {model_id} (Modeller) overlaid on 1HVY chain A"
    desc = ("Phase 6: Modeller AutoModel homology model "
            f"<code>{escape(model_pdb_path.name)}</code> aligned to the cleaned 1HVY "
            "crystal chain A. Magenta = crystal, green = model.")
    scores_html = (
        f"<b>Scores:</b> molpdf=<code>{scores.get('molpdf','?')}</code> · "
        f"DOPE=<code>{scores.get('DOPE','?')}</code> · "
        f"GA341=<code>{scores.get('GA341','?')}</code> · "
        f"Cα RMSD vs crystal=<code>{scores.get('rmsd_to_crystal','?')}</code> Å"
    )
    crystal_text = crystal_pdb_path.read_text().replace("`", "'")
    model_text = model_pdb_path.read_text().replace("`", "'")
    return VIEWER_TEMPLATE.format(
        title=title,
        description=desc,
        scores_html=scores_html,
        crystal_pdb=crystal_text,
        model_pdb=model_text,
    )


def update_viewers_index(viewer_files: list[str]) -> None:
    """Insert a 'Homology models' section into viewers/index.html, idempotent."""
    idx_path = VIEWERS_DIR / "index.html"
    html = idx_path.read_text()
    marker_open = "<!-- HOMOLOGY MODELS SECTION START -->"
    marker_close = "<!-- HOMOLOGY MODELS SECTION END -->"
    section = [marker_open]
    section.append('<section><h2>Homology models (Phase 6, Modeller)</h2><ul>')
    for vf in sorted(viewer_files):
        # vf is e.g. modeller_model01.html
        m_id = vf.replace("modeller_model", "").replace(".html", "")
        section.append(f'<li><a href="{vf}">Homology model {int(m_id)} '
                       'overlaid on 1HVY chain A</a></li>')
    section.append("</ul></section>")
    section.append(marker_close)
    new_section = "\n".join(section)

    if marker_open in html and marker_close in html:
        # Replace existing block
        before, _, rest = html.partition(marker_open)
        _, _, after = rest.partition(marker_close)
        html = before + new_section + after
    else:
        # Insert just before the trailing "</body>"
        insert_at = html.rfind("</body>")
        html = html[:insert_at] + new_section + "\n" + html[insert_at:]
    idx_path.write_text(html)
    LOG.info("Updated %s with %d homology-model links", idx_path, len(viewer_files))


def write_readme() -> Path:
    md = PHASE_DIR / "README_PHASE6.md"
    sel = json.loads((STEP2_DIR / "selected_templates.json").read_text())
    selected = sel["selected"]
    sel_lines = "\n".join(
        f"  - **{t['pdb_id']}_{t['chain']}** — {t['identity_pct']}% identity, "
        f"resolution {t.get('resolution','?')} Å"
        for t in selected
    )
    md.write_text(
f"""# Phase 6 — Homology modelling of human TYMS (UniProt P04818)

This phase builds a Modeller-based homology-modelling sub-pipeline. It is *additive*
to v1–v5: nothing in those folders is touched.

## (a) Why homology modelling?

Even when an experimental crystal structure exists for a target (here, 1HVY),
homology modelling is a useful *educational* exercise: it shows how a structure
is reconstructed from sequence + templates and lets you compare the model against
the known crystal as a ground-truth benchmark.

## (b) Why <100% identity templates?

We deliberately *excluded* 1HVY itself from BLAST hits and filtered hits to
**30% ≤ identity ≤ 95%**, **coverage ≥ 80%**, **resolution ≤ 2.5 Å**. This forces
Modeller to do real homology modelling rather than copying coordinates from the
identical-sequence template. In production one would obviously use the highest
quality template available, including 100% identity matches if they exist.

Selected templates:
{sel_lines}

## (c) What each step produced

1. `01_clean_pdb/` — Chain-A-only 1HVY (cleaned of HETATMs and chain B) +
   single-chain FASTA, sanity-checked against P04818.
2. `02_blast/` — Remote NCBI BLAST against the **pdb** database; filtered hits;
   downloaded selected template PDBs and extracted the relevant chains.
3. `03_alignment/` — Combined target + 3 templates with ClustalW; converted to
   Modeller PIR (`alignment.ali` / `alignment.pir`).
4. `04_modeller_run/` — 10 candidate models via Modeller AutoModel, scored with
   DOPE / molpdf / GA341 (`scores.csv`).
5. `05_comparison/` — PyMOL pairwise overlays + per-residue Cα-distance bar
   plots + all-models overlay + `rmsd_per_model.csv`.
6. `06_validation/` — Local Ramachandran φ/ψ plots + per-residue normalized DOPE
   profiles + `quality_overview.png` + `SAVES_MANUAL.md` for the UCLA SAVES
   step (manual upload).
7. `07_viewers/` — 3Dmol viewers per model (also mirrored into `viewers/` for
   GitHub Pages).

## (d) Best-model selection

Two complementary criteria:
- **Best by DOPE** (most negative ⇒ best fold-energy): see `04_modeller_run/best_by_dope.json`.
- **Best by Cα RMSD to crystal**: see `05_comparison/best_summary.json`.

In a real prediction (no crystal) only DOPE / molpdf / GA341 / Ramachandran are
available, so **best_by_dope** is the canonical choice and is copied to
`models/best_model.pdb`.

## (e) Limitations

- All 10 models are extremely similar (RMSD differences < 0.05 Å, DOPE within
  ~500 units), so the "best" choice is largely cosmetic.
- TYMS dimer interface is *not* modelled (single-chain target only) — for
  enzymatic studies the dimer should be reconstructed afterwards.
- Local Ramachandran is a *basic* PROCHECK substitute; for publication-quality
  validation use UCLA SAVES (PROCHECK / ERRAT / VERIFY3D / WHATCHECK) — see
  `06_validation/SAVES_MANUAL.md`.
- Templates were *intentionally* sub-100%; using 1HVY itself would give a
  near-perfect rebuild but defeat the educational purpose of the phase.

Generated {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.
"""
    )
    return md


def write_summary_json(rama_rows: list[dict]) -> Path:
    sel = json.loads((STEP2_DIR / "selected_templates.json").read_text())
    templates = [
        {"pdb_id": t["pdb_id"], "chain": t["chain"],
         "identity_pct": t["identity_pct"], "resolution_A": t.get("resolution")}
        for t in sel["selected"]
    ]
    # Best by DOPE
    scores_csv = STEP4_DIR / "scores.csv"
    best_dope = None
    rows = []
    with open(scores_csv) as fh:
        rows = list(csv.DictReader(fh))
    if rows:
        b = min(rows, key=lambda r: float(r["DOPE"]))
        best_dope = {"model_id": int(b["model_id"]),
                     "model_pdb": b["model_pdb"],
                     "DOPE": float(b["DOPE"]),
                     "molpdf": float(b["molpdf"]),
                     "GA341": float(b["GA341"])}
    # Best by RMSD
    rmsd_csv = STEP5_DIR / "rmsd_per_model.csv"
    best_rmsd = None
    rmsd_rows = []
    with open(rmsd_csv) as fh:
        rmsd_rows = list(csv.DictReader(fh))
    if rmsd_rows:
        b = min(rmsd_rows, key=lambda r: float(r["rmsd_to_crystal"]))
        best_rmsd = {"model_id": int(b["model_id"]),
                     "model_pdb": b["model_pdb"],
                     "rmsd_to_crystal": float(b["rmsd_to_crystal"]),
                     "n_atoms_aligned": int(b["n_atoms_aligned"])}

    identities = [t["identity_pct"] for t in sel["selected"]]
    summary = {
        "phase": "6 — Modeller homology modelling",
        "target_uniprot": "P04818",
        "target_pdb_source": "1HVY chain A",
        "n_models_built": len(rows),
        "templates": templates,
        "template_identity_range_pct": {
            "min": min(identities) if identities else None,
            "max": max(identities) if identities else None,
        },
        "best_by_dope": best_dope,
        "best_by_rmsd_to_crystal": best_rmsd,
        "ramachandran_summary": [
            {"model_id": r["model_id"],
             "pct_favoured": r["pct_favoured"],
             "pct_allowed": r["pct_allowed"],
             "pct_outlier": r["pct_outlier"]}
            for r in rama_rows
        ],
        "saves_web_upload_status": "manual (see 06_validation/SAVES_MANUAL.md)",
        "generated": datetime.now().isoformat(timespec="seconds"),
    }
    out = PHASE_DIR / "summary.json"
    out.write_text(json.dumps(summary, indent=2))
    return out


def build_docx_report() -> Path:
    """Standalone DOCX report for Phase 6."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("Phase 6 — Modeller Homology Modelling of TYMS (P04818)", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "This standalone report covers the Phase 6 sub-pipeline added to the "
        "TYMS project. The original v5 FINAL report is unchanged. Phase 6 builds "
        "homology models from BLAST-discovered, sub-100% identity templates "
        "(deliberately, for educational purposes) and validates them locally."
    )

    doc.add_heading("Section 17.1 — Templates selected", level=2)
    sel = json.loads((STEP2_DIR / "selected_templates.json").read_text())
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "PDB"
    hdr[1].text = "Chain"
    hdr[2].text = "% Identity"
    hdr[3].text = "Resolution (Å)"
    for t in sel["selected"]:
        row = table.add_row().cells
        row[0].text = t["pdb_id"]
        row[1].text = t["chain"]
        row[2].text = f"{t['identity_pct']:.2f}"
        row[3].text = str(t.get("resolution", "?"))
    doc.add_paragraph(sel.get("note", ""))

    doc.add_heading("Section 17.2 — Modeller AutoModel scores", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "molpdf"
    hdr[2].text = "DOPE"
    hdr[3].text = "GA341"
    with open(STEP4_DIR / "scores.csv") as fh:
        for row in csv.DictReader(fh):
            r = table.add_row().cells
            r[0].text = row["model_id"]
            r[1].text = f'{float(row["molpdf"]):.2f}'
            r[2].text = f'{float(row["DOPE"]):.2f}'
            r[3].text = f'{float(row["GA341"]):.4f}'

    doc.add_heading("Section 17.3 — RMSD vs crystal (1HVY chain A)", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "Cα RMSD (Å)"
    hdr[2].text = "n atoms aligned"
    with open(STEP5_DIR / "rmsd_per_model.csv") as fh:
        for row in csv.DictReader(fh):
            r = table.add_row().cells
            r[0].text = row["model_id"]
            r[1].text = f'{float(row["rmsd_to_crystal"]):.3f}'
            r[2].text = row["n_atoms_aligned"]

    doc.add_heading("Section 17.4 — Best model selection", level=2)
    summary = json.loads((PHASE_DIR / "summary.json").read_text())
    bd = summary["best_by_dope"]
    br = summary["best_by_rmsd_to_crystal"]
    p = doc.add_paragraph()
    p.add_run(f"Best by DOPE: model {bd['model_id']} ({bd['model_pdb']}), "
              f"DOPE={bd['DOPE']:.2f}.\n")
    p.add_run(f"Best by Cα RMSD to crystal: model {br['model_id']} ({br['model_pdb']}), "
              f"RMSD={br['rmsd_to_crystal']:.3f} Å.\n")
    p.add_run("In a blind prediction (no crystal), best-by-DOPE is the canonical pick; "
              "best-by-RMSD is shown for benchmarking only.")

    doc.add_heading("Section 17.5 — Quality overview", level=2)
    overview = STEP6_DIR / "quality_overview.png"
    if overview.exists():
        doc.add_picture(str(overview), width=Inches(6.2))

    doc.add_heading("Section 17.6 — All-models overlay (PyMOL)", level=2)
    overlay = STEP5_DIR / "all_models_overlay.png"
    if overlay.exists():
        doc.add_picture(str(overlay), width=Inches(6.2))

    doc.add_heading("Section 17.7 — Per-model figures", level=2)
    for i in range(1, 11):
        doc.add_heading(f"Model {i}", level=3)
        for img in [
            STEP5_DIR / f"pairwise_model{i:02d}.png",
            STEP6_DIR / f"ramachandran_model{i:02d}.png",
            STEP6_DIR / f"dope_profile_model{i:02d}.png",
        ]:
            if img.exists():
                doc.add_picture(str(img), width=Inches(5.6))

    doc.add_heading("Section 17.8 — Validation note (UCLA SAVES)", level=2)
    saves_md = STEP6_DIR / "SAVES_MANUAL.md"
    if saves_md.exists():
        for line in saves_md.read_text().splitlines():
            if line.startswith("# "):
                continue
            doc.add_paragraph(line)

    doc.add_heading("Section 17.9 — Limitations", level=2)
    doc.add_paragraph(
        "All 10 models are extremely similar (Cα RMSD differences < 0.05 Å). "
        "Only chain A was modelled; the dimer interface is not reconstructed. "
        "Local Ramachandran is a basic substitute for PROCHECK; the canonical "
        "validation workflow uses UCLA SAVES (PROCHECK / ERRAT / VERIFY3D / "
        "WHATCHECK) which has no programmatic API and must be run manually."
    )

    out = PROJECT_DIR / "09e_report_v5" / "report_PHASE6.docx"
    doc.save(str(out))
    return out


def main() -> int:
    STEP7_DIR.mkdir(parents=True, exist_ok=True)
    crystal_pdb_path = STEP1_DIR / "1hvy_chainA.pdb"

    # Per-model viewer scores
    scores_by_id: dict[int, dict] = {}
    with open(STEP4_DIR / "scores.csv") as fh:
        for row in csv.DictReader(fh):
            scores_by_id[int(row["model_id"])] = row
    rmsd_by_id: dict[int, dict] = {}
    with open(STEP5_DIR / "rmsd_per_model.csv") as fh:
        for row in csv.DictReader(fh):
            rmsd_by_id[int(row["model_id"])] = row

    models = sorted(MODELS_DIR.glob("target.B99990*.pdb"))
    viewer_files: list[str] = []
    for i, m in enumerate(models, start=1):
        scores = {
            "molpdf": f'{float(scores_by_id[i]["molpdf"]):.2f}',
            "DOPE": f'{float(scores_by_id[i]["DOPE"]):.2f}',
            "GA341": f'{float(scores_by_id[i]["GA341"]):.4f}',
            "rmsd_to_crystal": f'{float(rmsd_by_id[i]["rmsd_to_crystal"]):.3f}',
        }
        html = build_viewer_html(m, crystal_pdb_path, i, scores)
        out_local = STEP7_DIR / f"model{i:02d}.html"
        out_local.write_text(html)
        out_mirror = VIEWERS_DIR / f"modeller_model{i:02d}.html"
        out_mirror.write_text(html)
        viewer_files.append(out_mirror.name)
        LOG.info("Wrote viewer %s and mirror %s", out_local.name, out_mirror.name)

    update_viewers_index(viewer_files)

    # Ramachandran rows for summary
    rama_rows = []
    with open(STEP6_DIR / "ramachandran_stats.csv") as fh:
        for row in csv.DictReader(fh):
            rama_rows.append(row)

    summary_path = write_summary_json(rama_rows)
    LOG.info("Wrote %s", summary_path)
    readme_path = write_readme()
    LOG.info("Wrote %s", readme_path)
    docx_path = build_docx_report()
    LOG.info("Wrote DOCX report: %s (%d bytes)", docx_path, docx_path.stat().st_size)

    LOG.info("STEP 7 OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
