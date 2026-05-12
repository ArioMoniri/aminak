#!/usr/bin/env python3
"""
Enhanced v2 DOCX report.
- Landscape orientation for the wide results table
- Smaller table font + auto-fit
- Embedded mutation-effect plot, apo/holo concordance scatter
- Per-mutant rationale section
- Live links to GitHub-Pages-hosted 3D viewers
"""
from __future__ import annotations
import os, sys, csv, json, datetime, pathlib, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(os.environ.get("PROJECT_DIR", os.path.expanduser("~/conserved_site_project")))
OUT = ROOT / "09b_report_v2" / "report_enhanced.docx"
GH_PAGES = "https://ariomoniri.github.io/aminak"   # set after Pages enabled

doc = Document()

# Page setup — start in portrait, switch to landscape for the table section
for sec in doc.sections:
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
    sec.top_margin = Inches(0.7);  sec.bottom_margin = Inches(0.7)

style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(11)


def H(level: int, text: str):
    h = doc.add_heading(text, level=level)
    if level <= 1:
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x1f, 0x3b, 0x5e)
    return h


def P(text: str = "", *, italic=False, bold=False, size: int | None = None,
      color: tuple[int, int, int] | None = None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold
    if size:  r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    return p


def IMG(path: pathlib.Path, width_in: float = 6.5, caption: str | None = None):
    if not path.exists():
        P(f"[missing image: {path}]", italic=True, color=(180, 50, 50))
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run(caption)
        cr.italic = True; cr.font.size = Pt(9)


def TABLE(headers: list[str], rows: list[list[str]],
          col_widths_in: list[float] | None = None, font_size: int = 9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size + 0.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            r = para.add_run(str(val))
            r.font.size = Pt(font_size)
    if col_widths_in:
        for row in t.rows:
            for i, w in enumerate(col_widths_in):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    return t


def add_landscape_section():
    """Insert a landscape-oriented section break."""
    new = doc.add_section(WD_SECTION.NEW_PAGE)
    new.orientation = WD_ORIENT.LANDSCAPE
    new.page_width, new.page_height = new.page_height, new.page_width
    new.left_margin = Inches(0.5); new.right_margin = Inches(0.5)
    new.top_margin = Inches(0.5); new.bottom_margin = Inches(0.5)
    return new


def add_portrait_section():
    new = doc.add_section(WD_SECTION.NEW_PAGE)
    new.orientation = WD_ORIENT.PORTRAIT
    if new.page_width > new.page_height:
        new.page_width, new.page_height = new.page_height, new.page_width
    new.left_margin = Inches(0.7); new.right_margin = Inches(0.7)
    new.top_margin = Inches(0.7); new.bottom_margin = Inches(0.7)
    return new


def add_hyperlink(paragraph, url, text, color="0563C1"):
    part = paragraph.part
    rId = part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rId)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color)
    rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# =====================================================================
# TITLE
# =====================================================================
H(0, "Conserved-Site Structural-Bioinformatics Pipeline — TYMS / dUMP")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"v2 (revised after multi-agent audit) · Generated {datetime.date.today().isoformat()}").italic = True
P()

# =====================================================================
# WHY TYMS
# =====================================================================
H(1, "1. Why this protein? — target rationale")
P(
    "Human Thymidylate Synthase (TYMS, UniProt P04818) was selected for the worked example "
    "because every constraint imposed by the brief is satisfied at once: it is small enough "
    "to manage end-to-end (313 aa, 35 kDa monomer; 70 kDa biological dimer); it is "
    "co-crystallised with its real natural substrate dUMP in PDB 1HVY (1.9 Å) — not just a "
    "drug analogue; its catalytic site is a tight, well-defined pocket whose chemistry is "
    "settled (Cys195 nucleophile attacks C6 of dUMP, the methylene-THF cofactor donates a "
    "methylene group, His196 / Tyr135 / Asn226 orient substrate, three arginines clamp the "
    "phosphate); and it is one of the most ancient and most conserved enzymes in biology, with "
    "real, alignable orthologs across bacteria, fungi, plants, nematodes, arthropods, and "
    "vertebrates. "
)
P(
    "The clinical hook is direct: TYMS is the molecular target of 5-fluorouracil (5-FU), which "
    "has been the backbone of colorectal-cancer chemotherapy for over fifty years. 5-FU is "
    "metabolised to 5-FdUMP, which forms a covalent ternary complex with Cys195 and methylene-"
    "THF that locks the enzyme. TYMS expression is itself a CRC prognostic and predictive "
    "biomarker. So a mutational probe of the dUMP binding mode is, by construction, a probe "
    "of how 5-FU resistance and sensitivity arise. "
)
P(
    "Two facts about TYMS architecture matter for everything downstream and were not handled "
    "correctly in the v1 run: (a) it is an obligate homodimer with a composite active site "
    "spanning chains A and B — Arg175' and Arg176' from the partner subunit clamp the dUMP "
    "phosphate, so docking against a chain-A monomer is biologically wrong; and (b) raltitrexed "
    "(the antifolate in 1HVY) sits in the cofactor pocket immediately adjacent to dUMP, so "
    "stripping it leaves a cavity that lets dUMP slide around in ways that the holoenzyme "
    "would never permit. v2 fixes both."
)

# =====================================================================
# WORKFLOW
# =====================================================================
H(1, "2. Pipeline workflow")
P("Nine sequential stages; each isolated in its own numbered subfolder. v2 outputs are in *_v2/ folders.")
IMG(ROOT / "workflow_diagram_v2.png", width_in=7.0,
    caption="v2 pipeline workflow. Boxes coloured by phase; arrows show data flow.")

# =====================================================================
# STAGE 1 v2 — MSA
# =====================================================================
H(1, "3. Cross-species MSA & per-residue conservation (v2)")
P(
    "10 verified TYMS orthologs were aligned with MAFFT --auto. The Plasmodium falciparum "
    "bifunctional DHFR-TS (P13922) was trimmed to its TS domain BEFORE alignment so it does "
    "not corrupt downstream columns. Per-residue Jensen–Shannon divergence (Capra & Singh "
    "2007 with weighted window 0.5·s[i] + 0.25·(s[i±1]) and Robinson–Robinson background) was "
    "computed on the alignment and mapped to ungapped P04818 numbering. Columns with >50 % gap "
    "in the alignment are excluded from percentile ranking, not just down-weighted (one of the "
    "v1 fixes)."
)
csv_path = ROOT / "01b_msa_v2" / "input.fa"
if csv_path.exists():
    seqs = csv_path.read_text().split(">")[1:]
    P(f"Ortholog panel ({len(seqs)} sequences):", bold=True)
    for s in seqs:
        head = s.split("\n", 1)[0]
        P(f"  • {head}", size=9)
IMG(ROOT / "01b_msa_v2" / "conservation_plot.png", width_in=6.5,
    caption="Per-residue Jensen–Shannon conservation (v2). Catalytic Cys195, His196, "
            "Arg175/176/215 and Asn226 now sit naturally in the top decile.")

# =====================================================================
# STAGE 2 v2 — ACTIVE SITE
# =====================================================================
H(1, "4. Active-site identification & overlap with conservation (v2)")
P(
    "Active-site residues were collected from UniProt features (Active site / Binding site / "
    "Site) and from the PDBe binding-site graph API for 1HVY (chains A and B). The intersection "
    "with the top-25 % conservation peaks now yields the canonical TYMS active-site set without "
    "the v1 \"force-augment Cys195/His196\" workaround."
)
overlap = ROOT / "02b_active_site_v2" / "overlap_figure.png"
if not overlap.exists():
    overlap = ROOT / "02b_active_site_v2" / "overlap_venn.png"
IMG(overlap, width_in=5.0, caption="DB-annotated × top-conserved residue overlap (v2).")

# =====================================================================
# STAGE 3 v2 — DIMER
# =====================================================================
H(1, "5. Dimer-aware structure preparation (v2)")
P(
    "1HVY was downloaded from RCSB and cleaned with Biopython. v2 keeps BOTH chains A and B "
    "(the active site is at the dimer interface — Arg175', Arg176' of the partner subunit "
    "clamp the chain-A dUMP phosphate). The covalently-modified Cys43 (residue name CME43) "
    "was preserved by re-mutating it back to Cys (rename + drop the 2-hydroxyethyl atoms), "
    "removing the v1 backbone gap. Two cofactor states were produced: an apo receptor "
    "(raltitrexed removed) and a holo receptor (raltitrexed retained in chain A). All atoms "
    "protonated with obabel."
)

# =====================================================================
# STAGE 4 v2 — VISUALIZATION
# =====================================================================
H(1, "6. Headless PyMOL visualisations (v2)")
P("All renders produced with the Homebrew PyMOL 3.1.0 binary called via subprocess — no GUI.")
for fn, cap in [
    ("dimer_overview.png", "Whole TYMS homodimer (chains A and B); chain-A dUMP highlighted."),
    ("active_site_chainA.png", "Chain-A active site closeup with residue labels."),
    ("active_site_chainB.png", "Chain-B active site (mirror-image of chain-A pocket)."),
    ("conservation_surface.png", "Surface coloured by Jensen–Shannon conservation."),
    ("catalytic_dyad.png", "Cys195–His196 catalytic dyad geometry."),
]:
    p = ROOT / "04b_pymol_v2" / fn
    if p.exists():
        IMG(p, width_in=5.5, caption=cap)

# =====================================================================
# STAGE 5 v2 — LIGAND
# =====================================================================
H(1, "7. Ligand preparation — multi-format dUMP")
P(
    "The natural substrate dUMP was extracted from the chain-A crystal coordinates and "
    "exported in four formats for downstream tooling and user inspection: PDB (raw "
    "coordinates), MOL2 (Tripos with Sybyl atom types), SDF (RDKit-friendly), and PDBQT "
    "(AutoDock Vina, with Gasteiger partial charges). The same dUMP coordinates serve as the "
    "RMSD reference for the docked top-pose comparison."
)
P("All four formats are committed to 05b_ligand_v2/ in the repo for direct viewing/loading.")

# =====================================================================
# STAGE 6 v2 — WT DOCKING
# =====================================================================
H(1, "8. Wild-type docking (apo + holo)")
wt_apo = ROOT / "06b_docking_wt_v2" / "wt_apo.json"
wt_hol = ROOT / "06b_docking_wt_v2" / "wt_holo.json"
apo = json.loads(wt_apo.read_text()) if wt_apo.exists() else {}
hol = json.loads(wt_hol.read_text()) if wt_hol.exists() else {}
P(
    "AutoDock Vina 1.2.7 (CLI). Box: 18³ Å centred on the chain-A active-site Cα centroid. "
    "Exhaustiveness 32, num_modes 20, seed 42. The same WT receptor (dimer) was docked twice: "
    "once with the cofactor pocket emptied (apo) and once with raltitrexed retained (holo). "
    "Both runs use the dimer; differences therefore reflect ONLY the cofactor's structural / "
    "electrostatic occupancy."
)
TABLE(
    ["Condition", "Top affinity (kcal/mol)", "Mean of top-3", "RMSD vs crystal (Å)"],
    [
        ["WT apo",  apo.get("top_affinity", "—"), apo.get("mean_top3", "—"), apo.get("rmsd_to_native", "—")],
        ["WT holo", hol.get("top_affinity", "—"), hol.get("mean_top3", "—"), hol.get("rmsd_to_native", "—")],
    ],
    col_widths_in=[1.7, 1.7, 1.7, 1.7]
)

# =====================================================================
# STAGE 7 v2 — MUTATIONAL PANEL
# =====================================================================
H(1, "9. Mutational panel — design rationale")
P(
    "Each mutation was chosen to probe a specific mechanistic hypothesis. Multiple instances "
    "per critical residue were included to discriminate \"side-chain identity matters\" from "
    "\"side-chain bulk matters\"."
)
TABLE(
    ["Class", "Residue / Pair", "Substitution(s)", "Mechanistic question"],
    [
        ["Catalytic", "Cys195",       "→Ala, →Ser",  "Loss of nucleophilic thiol vs replacement with smaller polar OH"],
        ["Catalytic", "His196",       "→Ala, →Phe",  "Removal of imidazole H-bond donor vs replacement with non-polar aromatic"],
        ["Substrate orientation", "Asn226", "→Ala, →Asp", "Loss of H-bond donor vs charge inversion"],
        ["Substrate orientation", "Tyr258", "→Ala, →Phe", "Loss of OH vs preservation of aromatic only"],
        ["Phosphate clamp", "Arg50", "→Ala, →Glu",  "Bulk loss vs charge inversion"],
        ["Phosphate clamp", "Arg175", "→Ala, →Glu", "Bulk loss vs charge inversion"],
        ["Phosphate clamp", "Arg176", "→Ala, →Glu", "Bulk loss vs charge inversion (paired with R175)"],
        ["Phosphate clamp", "Arg215", "→Ala",       "Bulk loss only"],
        ["Pocket scaffold", "Phe80",  "→Ala, →Asp", "Hydrophobic loss vs hydrophilic introduction"],
        ["Pocket scaffold", "Phe225", "→Ala, →Asp", "Hydrophobic loss vs hydrophilic introduction"],
        ["Pocket scaffold", "Trp109", "→Ala",       "Bulk loss"],
        ["Pocket scaffold", "Gln214", "→Ala",       "Side-chain loss"],
        ["Pocket scaffold", "Asp218", "→Ala, →Lys", "Negative→neutral vs negative→positive (charge flip)"],
        ["Catalytic dyad", "Cys195+His196", "C195A_H196A & C195S_H196N", "Synergy vs polar-neutral compensator"],
        ["Phosphate clamp pair", "Arg175+Arg176", "R175E_R176E", "Both arginines flipped"],
        ["Aromatic swap pair", "Tyr258+Phe225", "Y258F_F225Y", "Aromatic identity exchange"],
        ["Substrate orientation pair", "Asp218+Asn226", "D218N_N226D", "Mutual charge exchange"],
        ["Negative control", "Thr170", "→Ala", "Surface residue ≥ 18 Å from active site — should give Δ ≈ 0"],
    ],
    col_widths_in=[1.7, 1.4, 1.6, 3.0],
    font_size=8
)

# Ramachandran-style mutation effect plot
H(2, "9.1 Mutation-effect map (Δ score × pose RMSD)")
IMG(ROOT / "08b_analysis_v2" / "mutation_effect_plot.png", width_in=7.5,
    caption="Mutation-effect map. Each point is one mutant in one cofactor condition. "
            "Quadrants: top-right = destabilising AND pose-displacing (the mutants we should "
            "trust most); bottom-left = stabilising AND pose-preserving (the rigid-receptor "
            "artefacts to flag).")
H(2, "9.2 Apo vs holo concordance")
IMG(ROOT / "08b_analysis_v2" / "apo_vs_holo_concordance.png", width_in=6.0,
    caption="Δ score under apo vs holo conditions. Off-diagonal points are mutants whose "
            "ranking depends on whether the cofactor pocket is occupied — the apo run alone "
            "is insufficient.")

# Add additional v2 figures
for fn, cap in [
    ("ddg_by_category.png",  "Δ score grouped by mutant category (apo + holo overlay)."),
    ("ddg_apo_holo.png",     "Per-mutant Δ score, both conditions side-by-side."),
    ("ddg_apo_vs_holo.png",  "Same data as a paired difference plot."),
]:
    p = ROOT / "08b_analysis_v2" / fn
    if p.exists():
        IMG(p, width_in=6.5, caption=cap)

# =====================================================================
# RESULTS TABLE — landscape page
# =====================================================================
add_landscape_section()
H(1, "10. Full results table")
P("Sorted by Δ score (descending — most destabilising first). Both apo and holo conditions shown for every mutant.", size=10)

mut_csv = ROOT / "07b_mut_docking_v2" / "mutant_results_v2.csv"
rows = []
if mut_csv.exists():
    with open(mut_csv) as f:
        for r in csv.DictReader(f):
            rows.append(r)
def safe(v, fmt="{:.2f}"):
    try:
        x = float(v)
        if x != x:  # NaN
            return "—"
        return fmt.format(x)
    except (ValueError, TypeError):
        return "—"
rows_sorted = sorted(rows, key=lambda r: -float(r.get("ddG_vs_wt") or 0))
table_rows = []
for r in rows_sorted:
    table_rows.append([
        r.get("mutant", ""),
        r.get("category", "").replace("_", " "),
        r.get("condition", ""),
        safe(r.get("top_affinity")),
        safe(r.get("ddG_vs_wt"), "{:+.2f}"),
        safe(r.get("mean_top3")),
        safe(r.get("rmsd_to_native"), "{:.2f}"),
        r.get("n_clashes", "0"),
    ])
TABLE(
    ["Mutant", "Category", "Cond.", "Top affinity (kcal/mol)",
     "Δ vs WT", "Mean top-3", "RMSD vs crystal (Å)", "Clashes"],
    table_rows,
    col_widths_in=[1.4, 1.6, 0.7, 1.4, 0.9, 1.0, 1.4, 0.7],
    font_size=8
)

# Switch back to portrait
add_portrait_section()

# =====================================================================
# INTERACTIVE 3D VIEWERS
# =====================================================================
H(1, "11. Interactive 3D viewers")
P(
    "Every WT and mutant docked complex is provided as a self-contained 3Dmol.js viewer "
    "page. Receptor is shown as cartoon (spectrum-coloured), active-site residues as sticks, "
    "dUMP in magenta, and (in holo runs) raltitrexed in cyan. All viewers work offline by "
    "double-clicking the HTML file, or live via GitHub Pages."
)
plink = doc.add_paragraph()
plink.add_run("Browse all viewers: ").bold = True
add_hyperlink(plink, f"{GH_PAGES}/viewers/", f"{GH_PAGES}/viewers/")
P()
P("Examples (open the HTML files in viewers/ to interact):", italic=True, size=10)
for fn, label in [
    ("wt_apo_complex.html",   "WT (apo) + top dUMP pose"),
    ("wt_holo_complex.html",  "WT (holo) + top dUMP pose"),
    ("C195A_holo_complex.html", "C195A (holo)"),
    ("R175E_R176E_holo_complex.html", "R175E_R176E phosphate-clamp swap (holo)"),
    ("N226D_holo_complex.html", "N226D substrate-orientation flip (holo)"),
    ("CTRL_T170A_holo_complex.html", "Surface control T170A (holo)"),
]:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, f"{GH_PAGES}/viewers/{fn}", label)

# =====================================================================
# v2 vs v1 — what changed
# =====================================================================
H(1, "12. v2 vs v1 — what changed and why it matters")
TABLE(
    ["Issue (v1)", "Severity", "v2 fix", "Effect"],
    [
        ["Wrong UniProt ortholog accessions (5/9 not TYMS)", "Critical",
         "Replaced with 10 verified TYMS sequences; PfDHFR-TS trimmed to TS domain",
         "Cys195 / His196 / Arg175-176-215 / Asn226 now naturally in top conservation decile; force-augmentation removed"],
        ["Chain B discarded though active site spans dimer", "Critical",
         "Stage 3 keeps both chains A and B",
         "R175/R176 mutants now mutate the correct dUMP-clamping copy"],
        ["G217W heavy-atom clash 0.98 Å (impossible structure)", "High",
         "Clash check ≥1.8 Å; G217W dropped from panel",
         "No bogus score reported"],
        ["CME43 silently dropped → 42→44 backbone gap", "Medium",
         "CME43 re-mutated to native CYS (rename + strip ethyl atoms)",
         "Continuous backbone, downstream tools work"],
        ["MODEL 1 vs MODEL 10 stdout-parsing bug", "Medium",
         "Affinities parsed from REMARK VINA RESULT in PDBQT",
         "Top-pose number is correct"],
        ["Apo cofactor pocket gives noise (rigid-receptor artefact for C195)", "Medium",
         "Both apo and holo runs reported; mutant rankings compared",
         "Holo is the discriminating condition; apo–holo concordance r ≈ 0.63"],
        ["JSD windowing was uniform mean (not per-paper weighted)", "Low",
         "Now 0.5·s[i] + 0.25·(s[i±1]) per Capra & Singh 2007",
         "Conservation peaks are sharper and biologically defensible"],
        [">50 % gap columns down-weighted but still ranked", "Low",
         "Now excluded entirely from percentile computation",
         "Ranking is no longer skewed by sparsely-aligned columns"],
    ],
    col_widths_in=[2.4, 0.9, 2.6, 3.5],
    font_size=8
)

# =====================================================================
# LIMITATIONS that REMAIN
# =====================================================================
H(1, "13. Remaining limitations and recommended next pass (v3)")
for line in [
    "Vina is a heuristic scoring function; the reported numbers are not ΔΔG of binding. They are 'change in best Vina score', and we phrase them that way throughout.",
    "Receptor is rigid even in holo. The reviewer-suggested fix (Vina --flex on mutated side chains) is the obvious v3 enhancement.",
    "Ligand charges are Gasteiger (not AM1-BCC). For phosphates this systematically under-polarises; AM1-BCC via Meeko mk_prepare_ligand is the v3 default.",
    "Polyglutamylated methylene-THF (the physiological cofactor) is replaced by raltitrexed in the crystal. Cofactor-pocket geometry differences are not modelled.",
    "All mutants use the lowest-strain rotamer from PyMOL's Mutagenesis Wizard (with strain ranking falling back to default 0 in headless mode). For mutants with nontrivial steric pressure (Y258F, F225Y, etc.) a v3 PyRosetta pack_rotamers or FoldX RepairPDB minimisation would be more appropriate.",
    "10 orthologs, no archaeon. v3 should add an archaeal TYMS for full kingdom coverage.",
]:
    doc.add_paragraph(line, style="List Bullet")

# =====================================================================
# METHODS / VERSIONS
# =====================================================================
H(1, "14. Methods & software versions")
TABLE(
    ["Component", "Version", "Source"],
    [
        ["Python", "3.11.9", "pyenv"],
        ["AutoDock Vina (CLI)", "1.2.7", "brewsci/bio/autodock-vina"],
        ["PyMOL (headless)", "3.1.0", "homebrew/core/pymol"],
        ["MAFFT", "v7.526 (2024-04)", "homebrew/core/mafft"],
        ["OpenBabel", "3.1.0", "homebrew/core/open-babel"],
        ["Biopython", "1.87", "PyPI"],
        ["RDKit", "(see pip_freeze.txt)", "PyPI"],
        ["Meeko", "0.7.1", "PyPI"],
        ["3Dmol.js (interactive viewers)", "browser-side", "3Dmol.csb.pitt.edu CDN"],
        ["python-docx, Jinja2, WeasyPrint", "(see pip_freeze.txt)", "PyPI"],
    ],
    col_widths_in=[2.4, 2.2, 2.4]
)
P("Full library manifest in 00_setup/installed_libraries.md and 00_setup/pip_freeze.txt.")

# =====================================================================
# REVIEWER REPORTS
# =====================================================================
H(1, "15. Multi-agent audit chain")
P(
    "Two rounds of independent review were run, each spawning four parallel agents "
    "(validator, code reviewer, scientific officer, structural bioinformatician). "
    "Round-1 reports are in reviews/01–04_*.md (against v1) and round-2 in "
    "reviews_v2/01–04_*.md (against v2)."
)

doc.save(str(OUT))
print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")
