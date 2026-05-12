#!/usr/bin/env python3
"""Stage 9 v3: HTML/PDF/DOCX report with limitations, Delta Vina wording, echoed params."""
import os, sys, base64, json, math
from datetime import datetime
import pandas as pd
from jinja2 import Template

PROJECT = os.path.expanduser("~/conserved_site_project")
REP_DIR = os.path.join(PROJECT, "09c_report_v3")
LOG_DIR = os.path.join(PROJECT, "logs")
LOG_FILE = os.path.join(PROJECT, "pipeline.log")
STAGE_LOG = os.path.join(LOG_DIR, "v3_09_report.log")

MSA_DIR = os.path.join(PROJECT, "01b_msa_v2")
AS_DIR = os.path.join(PROJECT, "02b_active_site_v2")
PYM_DIR = os.path.join(PROJECT, "04b_pymol_v2")
WT_DIR = os.path.join(PROJECT, "06c_docking_wt_v3")
MUT_DIR = os.path.join(PROJECT, "07c_mut_docking_v3")
ANA_DIR = os.path.join(PROJECT, "08c_analysis_v3")


def log(msg):
    line = f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [V3] STAGE9: {msg}"
    print(line, flush=True)
    with open(LOG_FILE, "a") as f:
        f.write(line + "\n")
    with open(STAGE_LOG, "a") as f:
        f.write(line + "\n")


def b64_img(path):
    if not os.path.exists(path):
        return ""
    with open(path, "rb") as f:
        return "data:image/png;base64," + base64.b64encode(f.read()).decode("ascii")


TEMPLATE = """<!doctype html>
<html><head><meta charset="utf-8"><title>TYMS Conserved Site v3 Report</title>
<style>
@page { size: A4; margin: 18mm; }
body { font-family: -apple-system, "Helvetica Neue", Helvetica, Arial, sans-serif;
       color: #222; line-height: 1.45; font-size: 10.5pt; }
h1 { color: #1a3759; border-bottom: 2px solid #1a3759; padding-bottom: 6px; }
h2 { color: #1a3759; margin-top: 22px; border-bottom: 1px solid #aac; padding-bottom: 3px; }
h3 { color: #2c4a6f; margin-top: 14px; }
img { max-width: 100%; height: auto; border: 1px solid #ccc; margin: 6px 0; }
table { border-collapse: collapse; font-size: 9pt; margin: 8px 0; width: 100%; }
th, td { border: 1px solid #888; padding: 4px 8px; text-align: left; }
th { background: #eef2f7; }
.small { font-size: 9pt; color: #555; }
.kbd { font-family: Menlo, monospace; background: #f4f4f4; padding: 1px 4px; border-radius: 3px; }
.fig { text-align: center; margin: 10px 0; }
.cap { font-size: 9pt; color: #555; font-style: italic; }
.box { background: #f7f9fc; border-left: 4px solid #1a3759; padding: 8px 12px; margin: 10px 0; }
.warn { background: #fff7e6; border-left: 4px solid #d4a017; padding: 8px 12px; margin: 10px 0; }
.lim { background: #fdecea; border-left: 4px solid #c0392b; padding: 8px 12px; margin: 10px 0; }
</style></head><body>

<h1>TYMS conserved active site — pipeline v3 report</h1>
<div class="small">Generated {{ now }} • Project root: <span class="kbd">~/conserved_site_project</span> • Pipeline version 3</div>

<div class="box">
<b>v3 changes vs v2 (audit fixes):</b>
<ul>
<li><b>Receptor charges</b>: rebuilt every PDBQT with Gasteiger charges via <span class="kbd">obabel -xr -p 7.4 --partialcharge gasteiger</span>. v2 had all-zero charges (broken meeko fallback).</li>
<li><b>WT reference re-docked</b>: box centred on crystal dUMP centroid, exhaustiveness 96, num_modes 32, box 22³ Å, dual-seed (42 + 7) sanity check.</li>
<li><b>Sign convention</b>: <span class="kbd">delta_vina_vs_wt = top_aff_mut − top_aff_wt</span> — positive = destabilising (Vina more negative = better binding, so positive Δ means worse binding).</li>
<li><b>Rotamer relaxation</b>: replaced broken <span class="kbd">get_strain()</span> with PyMOL <span class="kbd">sculpt</span> (3 × 20 cycles on whole prot after each Mutagenesis Wizard apply).</li>
<li><b>mean_topk</b>: <span class="kbd">mean(affinities[:min(3, n_modes)])</span> — robust when fewer than 3 modes converge; <span class="kbd">n_modes</span> column added.</li>
<li><b>Cofactor</b>: re-protonated at pH 7.4 to give the carboxylate tail its physiological −1 charge.</li>
<li><b>Mis-dock annotation</b>: rows with RMSD &gt; 3 Å are flagged <span class="kbd">mis_docked=True</span> and excluded from top-destabiliser rankings.</li>
<li><b>Naming</b>: column <span class="kbd">delta_vina_vs_wt</span>, prose "Δ Vina score" — no longer mislabelled "ΔΔG" (Vina is empirical, not free energy).</li>
<li><b>Plot legend</b>: keys match emitted category names (ala_scan, opposite, arg_clamp, etc.).</li>
<li><b>Vina parameters</b> are echoed into <span class="kbd">wt_*.json</span>.</li>
</ul>
</div>

<h2>1. WT docking — apo vs holo (rebuilt receptors)</h2>
<table><tr><th>Condition</th><th>Top affinity (kcal/mol)</th><th>mean top-k</th>
<th>n modes</th><th>RMSD top-pose vs crystal dUMP (Å)</th><th>Best seed</th></tr>
<tr><td>WT apo</td>
  <td>{{ "%.2f"|format(wt_apo.top_affinity) }}</td>
  <td>{{ "%.2f"|format(wt_apo.mean_topk) }}</td>
  <td>{{ wt_apo.n_modes }}</td>
  <td>{{ "%.2f"|format(wt_apo.rmsd_top_to_native) }}</td>
  <td>{{ wt_apo.best_seed }}</td>
</tr>
<tr><td>WT holo</td>
  <td>{{ "%.2f"|format(wt_holo.top_affinity) }}</td>
  <td>{{ "%.2f"|format(wt_holo.mean_topk) }}</td>
  <td>{{ wt_holo.n_modes }}</td>
  <td>{{ "%.2f"|format(wt_holo.rmsd_top_to_native) }}</td>
  <td>{{ wt_holo.best_seed }}</td>
</tr>
</table>
<p class="small">Receptor preparation method: {{ rec_method }}. Centroid: ({{ "%.2f, %.2f, %.2f"|format(centroid[0], centroid[1], centroid[2]) }}) Å. Exhaustiveness 96, num_modes 32, box 22 Å.</p>

<h2>2. Mutant panel</h2>
<p>Panel composition: {{ panel_size }} mutants (no G217W; the v2 exploratory variant was dropped — heavy-atom clashes on insertion). Each mutant docked under apo and holo conditions = {{ panel_size * 2 }} runs total.</p>

<div class="fig"><img src="{{ p_bar }}" alt="bar"/>
<div class="cap">Figure 1 — Δ Vina score vs WT for each mutant in apo (blue) and holo (orange). Grey bars are mis-docked (RMSD &gt; 3 Å, untrustworthy). Positive = destabilising.</div></div>

<div class="fig"><img src="{{ p_scatter }}" alt="scatter"/>
<div class="cap">Figure 2 — Apo vs holo Δ Vina concordance.</div></div>

<div class="fig"><img src="{{ p_cat }}" alt="cat"/>
<div class="cap">Figure 3 — Δ Vina by mutation category (well-docked rows only).</div></div>

<div class="fig"><img src="{{ p_eff }}" alt="eff"/>
<div class="cap">Figure 4 — Mutation-effect map: Δ Vina (x) vs pose RMSD (y), faded = mis-docked. Categories use corrected legend keys.</div></div>

<h3>Top 5 destabilising mutations — apo (well-docked only)</h3>
<table><tr><th>Mutant</th><th>category</th><th>Δ Vina</th><th>RMSD</th><th>n modes</th></tr>
{% for r in top_apo %}<tr>
  <td>{{ r.mutant }}</td><td>{{ r.category }}</td>
  <td>{{ "%+.2f"|format(r.delta_vina_vs_wt) }}</td>
  <td>{{ "%.2f"|format(r.rmsd_to_native) }}</td>
  <td>{{ r.n_modes }}</td>
</tr>{% endfor %}
</table>

<h3>Top 5 destabilising mutations — holo (well-docked only)</h3>
<table><tr><th>Mutant</th><th>category</th><th>Δ Vina</th><th>RMSD</th><th>n modes</th></tr>
{% for r in top_holo %}<tr>
  <td>{{ r.mutant }}</td><td>{{ r.category }}</td>
  <td>{{ "%+.2f"|format(r.delta_vina_vs_wt) }}</td>
  <td>{{ "%.2f"|format(r.rmsd_to_native) }}</td>
  <td>{{ r.n_modes }}</td>
</tr>{% endfor %}
</table>

<div class="lim">
<b>Limitations.</b>
<ul>
<li><b>Vina is empirical, not free energy.</b> Reported numbers are "Δ Vina score" (kcal/mol), not ΔΔG of binding. Treat as a relative ranking, not as a thermodynamic prediction.</li>
<li><b>Cofactor polyglutamylation</b> (the physiological state of methylene-THF) is not modelled. The bound raltitrexed is mono-glutamate.</li>
<li><b>Apo runs are dominated by mis-docking</b> in the empty cofactor pocket. They serve as a negative-control / contrast for the holo signal — not as biological readouts.</li>
<li><b>T170A surface control</b> behaviour is included as a sanity check (a residue not in the active site should give Δ ≈ 0).</li>
<li><b>Receptor and ligand are rigid.</b> PyMOL Mutagenesis Wizard rotamers are sculpt-relaxed but not relaxed against a relaxed pocket; large mutations (Y258F, F225Y, R175E, D218K-class swaps) may carry residual local strain.</li>
<li><b>Docking pose RMSD is computed against the crystal dUMP heavy atoms only</b> (no symmetry handling for ring atoms).</li>
</ul>
</div>

<h2>3. Reproducibility</h2>
<table>
<tr><th>Stage</th><th>Tool / parameters</th></tr>
<tr><td>Receptor prep</td><td>{{ rec_method }} • Gasteiger charges • pH 7.4 • max abs charge {{ "%.3f"|format(max_q) }}</td></tr>
<tr><td>Cofactor prep</td><td>obabel -h -p 7.4 (re-protonated for −1 carboxylate at pH 7.4)</td></tr>
<tr><td>WT docking</td><td>Vina 1.2.7 • exh=96 • num_modes=32 • box=22³Å • seeds 42 + 7 (best taken)</td></tr>
<tr><td>Mutant docking</td><td>Vina 1.2.7 • exh=32 • num_modes=20 • box=22³Å • seed=42</td></tr>
<tr><td>Mutagenesis</td><td>PyMOL 3.1.0 headless • Mutagenesis Wizard (frame=1) + sculpt 3×20 cycles</td></tr>
<tr><td>Sign convention</td><td><span class="kbd">delta_vina_vs_wt = top_aff_mut − top_aff_wt</span> (positive = destabilising)</td></tr>
<tr><td>mean_topk</td><td><span class="kbd">mean(affinities[:min(3, n_modes)])</span></td></tr>
<tr><td>mis_docked</td><td>RMSD top-pose vs crystal dUMP &gt; 3 Å</td></tr>
</table>

<p class="small">Outputs: <span class="kbd">06c_docking_wt_v3/</span>, <span class="kbd">07c_mut_docking_v3/</span>, <span class="kbd">08c_analysis_v3/</span>, <span class="kbd">09c_report_v3/</span>. v1 and v2 remain untouched.</p>

</body></html>
"""


def main():
    os.makedirs(REP_DIR, exist_ok=True)
    open(STAGE_LOG, "w").close()
    log("Stage 9 v3 starting")

    wt_apo = json.load(open(os.path.join(WT_DIR, "wt_apo.json")))
    wt_holo = json.load(open(os.path.join(WT_DIR, "wt_holo.json")))
    sum8 = json.load(open(os.path.join(ANA_DIR, "summary_v3.json")))
    sum7 = json.load(open(os.path.join(MUT_DIR, "summary_v3.json")))

    centroid = wt_apo["vina_params"]["centroid"]
    rec_method = wt_apo["vina_params"]["receptor_methods"]["apo"]
    # Read max charge from PDBQT
    rec_pdbqt = os.path.join(WT_DIR, "protein_dimer_apo.pdbqt")
    max_q = 0.0
    with open(rec_pdbqt) as f:
        for line in f:
            if line.startswith("ATOM"):
                try:
                    c = float(line[66:76].strip())
                    if abs(c) > max_q: max_q = abs(c)
                except Exception:
                    pass

    tpl = Template(TEMPLATE)
    html = tpl.render(
        now=datetime.now().strftime("%Y-%m-%d %H:%M"),
        wt_apo=wt_apo, wt_holo=wt_holo,
        rec_method=rec_method, centroid=centroid, max_q=max_q,
        panel_size=sum7["n_panel"],
        p_bar=b64_img(os.path.join(ANA_DIR, "delta_vina_apo_holo.png")),
        p_scatter=b64_img(os.path.join(ANA_DIR, "delta_vina_apo_vs_holo.png")),
        p_cat=b64_img(os.path.join(ANA_DIR, "delta_vina_by_category.png")),
        p_eff=b64_img(os.path.join(ANA_DIR, "mutation_effect_plot.png")),
        top_apo=sum8["top5_destab_apo_clean"],
        top_holo=sum8["top5_destab_holo_clean"],
    )
    html_path = os.path.join(REP_DIR, "report.html")
    with open(html_path, "w") as f:
        f.write(html)
    log(f"wrote {html_path}")

    # PDF
    try:
        from weasyprint import HTML
        pdf_path = os.path.join(REP_DIR, "report.pdf")
        HTML(string=html, base_url=REP_DIR).write_pdf(pdf_path)
        log(f"wrote {pdf_path} ({os.path.getsize(pdf_path)} bytes)")
    except Exception as e:
        log(f"weasyprint failed: {e}")
        try:
            import subprocess
            subprocess.run(["wkhtmltopdf", html_path,
                            os.path.join(REP_DIR, "report.pdf")],
                           check=True, capture_output=True)
            log("PDF via wkhtmltopdf fallback")
        except Exception as e2:
            log(f"all PDF backends failed: {e2}")

    # DOCX
    try:
        from docx import Document
        from docx.shared import Inches
        doc = Document()
        doc.add_heading("TYMS conserved active site — pipeline v3 report", 0)
        doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True

        doc.add_heading("Summary of v3 audit fixes", 1)
        for line in [
            "Receptor PDBQT rebuilt with Gasteiger charges (obabel -xr -p 7.4 --partialcharge gasteiger).",
            "WT apo+holo re-docked with crystal-dUMP centroid, exhaustiveness 96, num_modes 32, box 22, dual seed.",
            "Sign convention: delta_vina_vs_wt = top_aff_mut - top_aff_wt; positive = destabilising.",
            "Rotamer relaxation via PyMOL sculpt (3x20 cycles); the broken get_strain() loop was removed.",
            "mean_topk = mean(affinities[:min(3, n_modes)]); n_modes column added.",
            "Cofactor re-protonated at pH 7.4.",
            "RMSD > 3 A flagged mis_docked and excluded from top destabilisers.",
            "Column and prose use 'Delta Vina score' instead of 'ddG_vs_wt'.",
            "Plot legend keys match actual category names.",
            "Full Vina parameters echoed into wt_*.json.",
        ]:
            doc.add_paragraph(line, style="List Bullet")

        doc.add_heading("WT docking (rebuilt receptors)", 1)
        t = doc.add_table(rows=1, cols=6)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Condition", "Top aff", "mean topk", "n modes", "RMSD vs native", "best seed"]):
            hdr[i].text = h
        for cond, dat in [("WT apo", wt_apo), ("WT holo", wt_holo)]:
            row = t.add_row().cells
            row[0].text = cond
            row[1].text = f"{dat['top_affinity']:.2f}"
            row[2].text = f"{dat['mean_topk']:.2f}"
            row[3].text = str(dat["n_modes"])
            row[4].text = f"{dat['rmsd_top_to_native']:.2f}"
            row[5].text = str(dat["best_seed"])

        doc.add_heading("Mutant panel — top 5 destabilising (apo, well-docked)", 1)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Mutant", "Category", "Delta Vina", "RMSD"]):
            hdr[i].text = h
        for r in sum8["top5_destab_apo_clean"]:
            row = t.add_row().cells
            row[0].text = r["mutant"]
            row[1].text = r["category"]
            row[2].text = f"{float(r['delta_vina_vs_wt']):+.2f}"
            row[3].text = f"{float(r['rmsd_to_native']):.2f}"

        doc.add_heading("Mutant panel — top 5 destabilising (holo, well-docked)", 1)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Mutant", "Category", "Delta Vina", "RMSD"]):
            hdr[i].text = h
        for r in sum8["top5_destab_holo_clean"]:
            row = t.add_row().cells
            row[0].text = r["mutant"]
            row[1].text = r["category"]
            row[2].text = f"{float(r['delta_vina_vs_wt']):+.2f}"
            row[3].text = f"{float(r['rmsd_to_native']):.2f}"

        for png_name, cap in [
            ("delta_vina_apo_holo.png", "Figure 1 - Delta Vina apo vs holo bar chart"),
            ("delta_vina_apo_vs_holo.png", "Figure 2 - Apo-vs-holo concordance"),
            ("delta_vina_by_category.png", "Figure 3 - Delta Vina by category (well-docked only)"),
            ("mutation_effect_plot.png", "Figure 4 - Mutation-effect map; faded = mis-docked"),
        ]:
            p = os.path.join(ANA_DIR, png_name)
            if os.path.exists(p):
                doc.add_picture(p, width=Inches(6.0))
                doc.add_paragraph(cap).italic = True

        doc.add_heading("Limitations", 1)
        for line in [
            "Vina is an empirical scoring function, not a free energy. Numbers reported are Delta Vina score, not Delta-Delta G of binding.",
            "Cofactor polyglutamylation (physiological state of methylene-THF) is not modelled.",
            "Apo runs are dominated by mis-docking in the empty cofactor pocket; they serve as a negative control / contrast for the holo signal, not as biological readouts.",
            "T170A surface control behaviour is included as a sanity check.",
            "Receptor and ligand are rigid. PyMOL Mutagenesis Wizard rotamers are sculpt-relaxed but not relaxed against a relaxed pocket.",
            "RMSD is computed against crystal dUMP heavy atoms only (no symmetry handling for ring atoms).",
        ]:
            doc.add_paragraph(line, style="List Bullet")

        doc.add_heading("Reproducibility", 1)
        doc.add_paragraph(
            f"Vina 1.2.7 (WT: exh=96, num_modes=32, box=22, seeds 42 + 7; mutants: exh=32, num_modes=20, box=22, seed=42). "
            f"Receptor: {rec_method}, Gasteiger charges, max |q| = {max_q:.3f}. "
            f"PyMOL 3.1.0 headless with Mutagenesis Wizard + sculpt. "
            f"Open Babel 3.x for hydrogens / PDBQT / Gasteiger charges. "
            f"Centroid (crystal dUMP, A): ({centroid[0]:.2f}, {centroid[1]:.2f}, {centroid[2]:.2f})."
        )
        doc.add_paragraph(f"All v3 outputs in 06c_docking_wt_v3, 07c_mut_docking_v3, "
                          f"08c_analysis_v3, 09c_report_v3. v1 and v2 left untouched.")

        docx_path = os.path.join(REP_DIR, "report.docx")
        doc.save(docx_path)
        log(f"wrote {docx_path} ({os.path.getsize(docx_path)} bytes)")
    except Exception as e:
        log(f"docx failed: {e}")
        import traceback
        log(traceback.format_exc())

    log("Stage 9 v3 DONE")


if __name__ == "__main__":
    main()
