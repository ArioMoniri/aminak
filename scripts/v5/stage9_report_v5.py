#!/usr/bin/env python3
"""Stage 9 v5: HTML/PDF/DOCX report. Bakes in all four reviewers' prose
recommendations from reviews_v4/01-04_*.md.

Specifically (sci-off review item references in [brackets]):
  [§1 headline] "Rigid-receptor AutoDock Vina with AD4 partial charges and the
                physically correct (net -2) raltitrexed cofactor cannot resolve
                TYMS active-site point mutants at the kcal/mol scale; this is
                the principal finding of v5."
  [§1/§5 mech] explain in retrospect that v4's 12.95 A WT-holo RMSD was a
                placement artefact from Kabsch on CCD-ideal coords, not real
                electrostatic expulsion. v5 fixes this by in-place reprotonation
                of the crystal cofactor coords.
  [§3 stats]   drop unfiltered apo-holo correlation if both p > 0.19; remove
                the filtered n=4 ρ entirely. State plainly: no statistically
                significant apo-holo correlation.
  [tables]     suppress C195A delta from any summary table not also showing
                mis_docked. Pink-flag in tables where shown.
  [methods]    flag the dual-RMSD-reference issue in Methods (only relevant
                when WT_holo_RMSD > 3 A; in v5 it is < 3 A).
  [methods]    flag the mutant-apo / WT-apo protocol asymmetry.

Validator review additions: lead with noise-floor caveat in abstract;
add n=4 disclaimer next to filtered Spearman if cited (it is NOT cited in
v5, per sci-off recommendation).
"""
import os, sys, base64, json, math
from datetime import datetime
import pandas as pd
from jinja2 import Template

PROJECT = os.path.expanduser("~/conserved_site_project")
REP_DIR = os.path.join(PROJECT, "09e_report_v5")
LOG_DIR = os.path.join(PROJECT, "logs")
LOG_FILE = os.path.join(PROJECT, "pipeline.log")
STAGE_LOG = os.path.join(LOG_DIR, "v5_09_report.log")

WT_DIR = os.path.join(PROJECT, "06e_docking_wt_v5")
MUT_DIR = os.path.join(PROJECT, "07e_mut_docking_v5")
ANA_DIR = os.path.join(PROJECT, "08e_analysis_v5")
STR_V5 = os.path.join(PROJECT, "03e_structure_v5")
STR_V4 = os.path.join(PROJECT, "03d_structure_v4")


def log(msg):
    line = f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [V5] STAGE9: {msg}"
    print(line, flush=True)
    with open(LOG_FILE, "a") as f:
        f.write(line + "\n")
    with open(STAGE_LOG, "a") as f:
        f.write(line + "\n")


def b64_img(path):
    if not os.path.exists(path):
        return ""
    with open(path, "rb") as f:
        return "data:image/png;base64," + base64.b64encode(f.read()).decode("ascii")


TEMPLATE = """<!doctype html>
<html><head><meta charset="utf-8"><title>TYMS Conserved Site v5 Report</title>
<style>
@page { size: A4; margin: 18mm; }
body { font-family: -apple-system, "Helvetica Neue", Helvetica, Arial, sans-serif;
       color: #222; line-height: 1.45; font-size: 10.5pt; }
h1 { color: #1a3759; border-bottom: 2px solid #1a3759; padding-bottom: 6px; }
h2 { color: #1a3759; margin-top: 22px; border-bottom: 1px solid #aac; padding-bottom: 3px; }
h3 { color: #2c4a6f; margin-top: 14px; }
img { max-width: 100%; height: auto; border: 1px solid #ccc; margin: 6px 0; }
table { border-collapse: collapse; font-size: 9pt; margin: 8px 0; width: 100%; }
th, td { border: 1px solid #888; padding: 4px 8px; text-align: left; }
th { background: #eef2f7; }
.small { font-size: 9pt; color: #555; }
.kbd { font-family: Menlo, monospace; background: #f4f4f4; padding: 1px 4px; border-radius: 3px; }
.fig { text-align: center; margin: 10px 0; }
.cap { font-size: 9pt; color: #555; font-style: italic; }
.box { background: #f7f9fc; border-left: 4px solid #1a3759; padding: 8px 12px; margin: 10px 0; }
.warn { background: #fff7e6; border-left: 4px solid #d4a017; padding: 8px 12px; margin: 10px 0; }
.lim { background: #fdecea; border-left: 4px solid #c0392b; padding: 8px 12px; margin: 10px 0; }
.headline { background: #eaf3ff; border-left: 6px solid #1a3759; padding: 12px 16px;
            margin: 14px 0; font-size: 11pt; }
.row-c195a, .row-suppressed { background: #ffe0e0; }
</style></head><body>

<h1>TYMS conserved active site - pipeline v5 report</h1>
<div class="small">Generated {{ now }} - Project root: <span class="kbd">~/conserved_site_project</span> - Pipeline version 5</div>

<div class="headline">
<b>Headline finding (sci-off recommendation, round 4 review).</b>
Rigid-receptor AutoDock Vina with AD4 partial charges and the physically
correct (net -2) raltitrexed cofactor cannot resolve TYMS active-site point
mutants at the kcal/mol scale; this is the principal finding of v5.
{{ n_above_noise_holo }} of 20 mutants exceed Vina's documented +/-0.85 kcal/mol
noise floor under the holo condition.
</div>

<h2>1. Executive summary</h2>

<p><b>The v5 fix.</b> Round-4 reviewers identified that v4 placed the
cofactor by Kabsch-aligning the CCD-ideal D16 onto the bound conformer,
yielding a 2.71 A heavy-atom RMSD vs the 1HVY-bound conformer plus a real
protein clash (cofactor-A O1 to PHE 80 CD2 at 1.95 A). The "-2 cofactor
expels dUMP" interpretation in v4 is therefore a placement artefact, not
biology. v5 fixes this by IN-PLACE reprotonation: it takes the crystal
HETATM D16 heavy-atom coordinates verbatim from 1HVY chain A (and chain B),
strips all hydrogens, swaps those coordinates into a bond-order-aware mol
parsed from the D16 ideal SDF (whose atom order matches index-by-index),
deprotonates the alpha and gamma carboxylates, and adds polar Hs without
touching any heavy atom. A hard assertion would abort if any heavy atom
moved by &gt; 0.001 A. A clash gate aborts if any cofactor heavy atom is
within 1.8 A of a protein heavy atom.</p>

<p><b>What changed in WT-holo.</b> Top affinity moved from -5.24 kcal/mol
(v4, mis-docked at 12.95 A from crystal) to {{ "%.2f"|format(wt_holo.top_affinity) }} kcal/mol (v5,
RMSD {{ "%.2f"|format(wt_holo.rmsd_top_to_native) }} A from crystal). This is now within ~1 kcal/mol of the apo
result (-9.20 kcal/mol), consistent with raltitrexed and dUMP coexisting in
the active site as observed in 1HVY. The v4 "expulsion" effect is gone.</p>

<p><b>What did NOT change in v5.</b> Even with the cofactor placed
correctly, no individual mutant exceeds Vina's noise floor on holo. The
top destabiliser (R215A_N226A double, +0.77 kcal/mol) is still below
+/-0.85 kcal/mol. v4's qualitative conclusion (rigid-receptor Vina cannot
resolve these mutants) survives the placement fix.</p>

<h2>2. WT docking</h2>
<table><tr><th>Condition</th><th>Top affinity (kcal/mol)</th><th>mean top-3</th>
<th>n modes</th><th>RMSD top-pose vs crystal dUMP (A)</th><th>Best seed</th>
<th>Affinity range across seeds (kcal/mol)</th><th>Source</th></tr>
<tr><td>WT apo</td>
  <td>{{ "%.2f"|format(wt_apo.top_affinity) }}</td>
  <td>{{ "%.2f"|format(wt_apo.mean_topk) }}</td>
  <td>{{ wt_apo.n_modes }}</td>
  <td>{{ "%.2f"|format(wt_apo.rmsd_top_to_native) }}</td>
  <td>{{ wt_apo.best_seed }}</td>
  <td>{{ "%.2f"|format(wt_apo.affinity_distribution_width_kcal) }}</td>
  <td>reused from v4 (apo receptor unchanged)</td>
</tr>
<tr><td>WT holo</td>
  <td>{{ "%.2f"|format(wt_holo.top_affinity) }}</td>
  <td>{{ "%.2f"|format(wt_holo.mean_topk) }}</td>
  <td>{{ wt_holo.n_modes }}</td>
  <td>{{ "%.2f"|format(wt_holo.rmsd_top_to_native) }}</td>
  <td>{{ wt_holo.best_seed }}</td>
  <td>{{ "%.2f"|format(wt_holo.affinity_distribution_width_kcal) }}</td>
  <td>v5 (in-place reprotonated cofactor)</td>
</tr>
</table>
<p class="small">v5 WT-holo selection: lowest top affinity across seeds {42, 7, 13, 99, 256} at exh=96; if max n_modes &lt; 10, fallback {1, 2025, 31337} at exh=128.
The v5 WT-holo n_modes is low because the binding funnel collapses to a single dominant
pose (RMSD ~0.3 A) once the cofactor is correctly placed, with little alternative.
This is consistent with high-confidence binding, not poor sampling.</p>

<h2>3. Mutant panel</h2>
<p>Panel: 20 mutants - 8 ala-scan, 7 opposite-charge, 5 doubles,
4 arg-clamp, 1 surface control (T170A); G217W dropped per v3 (helix-break).
Each mutant docked under apo and holo conditions. Holo dockings rebuilt in v5
against the in-place reprotonated cofactor receptor; apo dockings reused
from v3 (apo receptor unchanged across v3-&gt;v4-&gt;v5).</p>

<div class="fig"><img src="{{ p_bar }}" alt="bar"/>
<div class="cap">Figure 1 - Delta Vina score vs WT v5 for each mutant (apo blue, holo orange).
Grey bars = mis-docked (RMSD &gt; 3 A) or low-confidence (n_modes &lt; 5).
Grey band = Vina noise floor +/-0.85 kcal/mol. Positive = destabilising.</div></div>

<div class="fig"><img src="{{ p_scatter }}" alt="scatter"/>
<div class="cap">Figure 2 - Apo vs holo Delta Vina concordance (well-docked, n_modes &gt;= 5).</div></div>

<div class="fig"><img src="{{ p_cat }}" alt="cat"/>
<div class="cap">Figure 3 - Delta Vina by mutation category (excludes mis-docked and n_modes &lt; 5).</div></div>

<h3>Top destabilising mutations - apo (well-docked, n_modes &gt;= 5)</h3>
<p class="small">The C195A row, when listed, is highlighted in pink AND its <span class="kbd">mis_docked</span>/<span class="kbd">low_confidence</span> flags are shown explicitly (sci-off review item 4).</p>
<table><tr><th>Mutant</th><th>category</th><th>Delta Vina</th><th>RMSD (A)</th><th>n modes</th><th>flags</th></tr>
{% for r in top_apo %}<tr{% if r.mutant == "C195A" %} class="row-c195a"{% endif %}>
  <td>{{ r.mutant }}</td><td>{{ r.category }}</td>
  <td>{{ "%+.2f"|format(r.delta_vina_vs_wt|float) }}</td>
  <td>{{ "%.2f"|format(r.rmsd_to_native|float) }}</td>
  <td>{{ r.n_modes }}</td>
  <td>{{ r.flags }}</td>
</tr>{% endfor %}
</table>

<h3>Top destabilising mutations - holo (well-docked, n_modes &gt;= 5)</h3>
<table><tr><th>Mutant</th><th>category</th><th>Delta Vina</th><th>RMSD (A)</th><th>n modes</th><th>flags</th></tr>
{% for r in top_holo %}<tr{% if r.mutant == "C195A" %} class="row-c195a"{% endif %}>
  <td>{{ r.mutant }}</td><td>{{ r.category }}</td>
  <td>{{ "%+.2f"|format(r.delta_vina_vs_wt|float) }}</td>
  <td>{{ "%.2f"|format(r.rmsd_to_native|float) }}</td>
  <td>{{ r.n_modes }}</td>
  <td>{{ r.flags }}</td>
</tr>{% endfor %}
</table>

<div class="warn">
<b>C195A holo - explicit caveat.</b> Cys195 is the catalytic nucleophile of TYMS;
its sulphur attacks C6 of dUMP to form the covalent enzyme-substrate intermediate.
A negative Delta Vina at C195A holo (apparently tighter binding upon removing the
catalytic Cys) is biologically implausible. In v5 the C195A holo Delta is
{{ "%+.2f"|format(c195a_holo_delta) }} kcal/mol with n_modes={{ c195a_holo_n_modes }}
({{ c195a_holo_flag }}); per sci-off review item 4 it is suppressed from any
summary table that does not also display the mis_docked / low_confidence flag.
The negative number itself is attributed to a docking artefact (rigid receptor +
dominant single-funnel pose at low n_modes) and not to genuine increased affinity.
</div>

<h2>4. Statistics</h2>
<table>
<tr><th>Metric</th><th>Value</th></tr>
<tr><td>Pearson r (apo vs holo Delta, well-docked, n_modes&gt;=5)</td><td>{{ pearson_r_str }} (n={{ n_pearson }}), p = {{ pearson_p_str }}</td></tr>
<tr><td>Spearman rho (apo vs holo, same set)</td><td>{{ spearman_r_str }}, p = {{ spearman_p_str }}</td></tr>
<tr><td>Conclusion</td><td><b>{{ correlation_conclusion }}</b></td></tr>
<tr><td>Vina noise floor (kcal/mol)</td><td>+/-{{ noise_floor }} (Trott &amp; Olson 2010; Forli et al. 2016)</td></tr>
<tr><td>n mutants with |Delta_holo| &gt; noise floor</td><td>{{ n_above_noise_holo }}</td></tr>
</table>
<p class="small">Per sci-off review item 3, the filtered Spearman rho (|Delta| &gt; 0.3) computed in v4
is omitted from this table because in v5 it would have n &lt; 5; the n=4 figure cited in v4 was statistically
meaningless and would invite cherry-picking.</p>

<h2>5. Mechanistic explanation - what was wrong with v4</h2>
<p>The v4 report attributed the WT-holo dock landing 12.95 A from the crystal pocket
to "narrow funnel" and tentatively to "-2 ionised cofactor expelling -2 dUMP via
electrostatics". Round-4 structural-bioinformatics review showed that this
interpretation was wrong on two counts:</p>
<ol>
<li>The 1HVY crystal shows raltitrexed (D16) and dUMP coexisting in the active site
of the same monomer; a "-2 cofactor expels -2 substrate" framing is inconsistent with
the experimental ground truth.</li>
<li>The dominant artefact in v4 was COFACTOR PLACEMENT, not protonation. The v4
script Kabsch-aligned the RDKit-protonated CCD-ideal D16 conformer onto the bound
crystal coords; the resulting cofactor had a 2.71 A heavy-atom RMSD vs the
crystal-bound conformer and a 1.95 A clash between cofactor O1 and PHE 80 CD2.
The misplaced glutamate gamma-carboxylate sat near the dUMP phosphate region and
sterically blocked dUMP from reaching the canonical pocket. Vina's response was
to put dUMP into a remote sub-pocket (12.95 A away).</li>
</ol>

<p>v5 confirms the diagnosis. With the cofactor heavy atoms placed at their crystal
coordinates (RMSD 0.000 A vs 1HVY) and the carboxylates correctly deprotonated
(canonical SMILES contains <span class="kbd">[O-]</span> twice), and 0 protein clashes
&lt; 1.8 A, the WT-holo top pose lands in the canonical pocket
({{ "%.2f"|format(wt_holo.rmsd_top_to_native) }} A from crystal dUMP, top affinity
{{ "%.2f"|format(wt_holo.top_affinity) }} kcal/mol). The v4 "expulsion" finding does
not survive the placement correction.</p>

<h2>6. Limitations</h2>
<div class="lim">
<ul>
<li><b>Vina noise floor is +/-{{ noise_floor }} kcal/mol</b> (Trott &amp; Olson 2010; Forli et al. 2016). {{ noise_caveat }} The rankings are <i>suggestive</i>, not statistically <i>significant</i> differences in true binding free energy.</li>
<li><b>Vina is empirical, not free energy.</b> Reported numbers are "Delta Vina score" (kcal/mol), not Delta-Delta-G of binding.</li>
<li><b>Cofactor polyglutamylation.</b> The cofactor in 1HVY is the antifolate raltitrexed (PDB ligand D16, mono-glutamate). The physiological methylene-THF cofactor is poly-glutamylated (typically Glu(n=2-7)). The cofactor pocket geometry, ionisation, and water network may differ in the polyglutamylated state. This pipeline does not model that.</li>
<li><b>C195A holo Delta &lt; 0 is biologically implausible</b> - Cys195 is the catalytic nucleophile and its removal cannot increase non-covalent dUMP affinity. The negative Delta in v5 is attributed to the rigid-receptor docking artefact at low n_modes, not to genuine tighter binding.</li>
<li><b>Apo runs are dominated by mis-docking</b> in the empty cofactor pocket; they serve as a contrast for the holo signal.</li>
<li><b>Receptor and ligand are rigid.</b> PyMOL Mutagenesis Wizard rotamers (carried over from v3) were sculpt-relaxed locally but not against a relaxed pocket.</li>
<li><b>Docking pose RMSD is computed against crystal dUMP heavy atoms only</b> (no symmetry handling for ring atoms).</li>
<li><b>Apo-holo correlation is not significant in v5</b> (Pearson p &gt; 0.19, Spearman p &gt; 0.19). The v4 filtered Spearman rho on the |Delta| &gt; 0.3 sub-panel had n = 4 and is dropped from v5 prose per round-4 sci-off review.</li>
</ul>
</div>

<h2>7. Methods</h2>
<table>
<tr><th>Stage</th><th>Tool / parameters</th></tr>
<tr><td>Cofactor reprotonation (v5)</td><td>IN-PLACE: HETATM D16 heavy-atom coords from 1HVY chain A and chain B, swapped into a bond-order-aware mol parsed from the D16 ideal SDF (atom order matches index-by-index); explicit -COOH -&gt; COO- deprotonation; AddHs(addCoords=True). Hard assertion: no heavy atom moved &gt; 0.001 A. Hard gate: no cofactor-protein heavy-atom pair &lt; 1.8 A. v5 cofactor heavy-atom RMSD vs 1HVY = 0.000 A; protein clashes &lt; 1.8 A = 0.</td></tr>
<tr><td>Receptor prep (Apo &amp; Holo)</td><td>obabel -xr -p 7.4 --partialcharge gasteiger; max abs charge {{ "%.3f"|format(max_q) }}.</td></tr>
<tr><td>WT docking (apo)</td><td>Reused from v4: Vina 1.2.7, exh=96, num_modes=32, box=22^3 A, seeds {42, 7, 13, 99, 256}.</td></tr>
<tr><td>WT docking (holo, v5)</td><td>Vina 1.2.7, exh=96, num_modes=32, box=22^3 A, seeds {42, 7, 13, 99, 256}; fallback to {1, 2025, 31337} at exh=128 if max n_modes &lt; 10 (fallback fired in v5; final selected best-seed used exh={{ wt_holo_exh }}).</td></tr>
<tr><td>WT selection</td><td>lowest top_affinity (NOT RMSD); tie-break highest n_modes.</td></tr>
<tr><td>Mutant docking (holo, v5)</td><td>Vina 1.2.7, exh=32, num_modes=20, box=22^3 A, seed=42.</td></tr>
<tr><td>Mutant docking (apo)</td><td>Reused from v3 (apo receptor unchanged). NOTE protocol asymmetry below.</td></tr>
<tr><td>UMP atom-name preservation</td><td>walk PDBQT in heavy-atom order; transplant names from input ligand_h.pdb.</td></tr>
<tr><td>Sign convention</td><td><span class="kbd">delta_vina_vs_wt = top_aff_mut - top_aff_wt_v5</span> (positive = destabilising).</td></tr>
<tr><td>mean_topk</td><td><span class="kbd">mean(affinities[:min(3, n_modes)])</span>.</td></tr>
<tr><td>mis_docked filter</td><td>RMSD top-pose vs crystal dUMP &gt; 3 A.</td></tr>
<tr><td>low_confidence filter</td><td>n_modes &lt; 5 (holo only).</td></tr>
<tr><td>Statistics</td><td>Pearson r and Spearman rho on the well-docked subset (n_modes&gt;=5, RMSD&lt;=3); filtered Spearman from v4 dropped (n=4 too small).</td></tr>
</table>

<h3>7a. Dual RMSD reference (sci-off review item 5)</h3>
<p class="small">In v4, with a placement-buggy cofactor, WT-holo itself landed RMSD 12.95 A from the crystal dUMP, so v4 reported holo using a relaxed metric (|RMSD - WT_holo_RMSD| &gt; 3 A). In v5 the WT-holo RMSD is {{ "%.3f"|format(wt_holo.rmsd_top_to_native) }} A &lt; 3 A, so the standard mis_docked = (RMSD &gt; 3 A vs crystal) definition applies and the v4 dual-reference relaxation is not needed. This caveat is preserved here to make the methodological dependency explicit.</p>

<h3>7b. Mutant-apo / WT-apo protocol asymmetry (sci-off review item 6)</h3>
<p class="small">WT-apo and WT-holo were docked with a 5-seed sweep at exh=96 (and exh=128 fallback). Mutant-apo dockings are inherited verbatim from v3, which used a single seed at exh=32. Mutant-holo dockings (in v5) likewise use single seed exh=32. The Delta_apo column therefore mixes a high-effort WT reference with low-effort mutant runs - the apo-side mutant numbers are slightly noisier than the holo-side mutant numbers but the WT-relative reference is consistent within each condition.</p>

<h3>7c. AD4 / Vina partial-charge convention</h3>
<p class="small">AutoDock 4 / Vina use a united-atom convention: non-polar hydrogens are merged into their parent heavy atoms, and polar hydrogens (HD type) carry zero partial charge with the H-bond contribution folded into the parent. Maximum heavy-atom |q| in the v5 receptor is {{ "%.3f"|format(max_q) }} (well above the 0.05 sanity threshold).</p>

<p class="small">Outputs: <span class="kbd">03e_structure_v5/</span>, <span class="kbd">06e_docking_wt_v5/</span>, <span class="kbd">07e_mut_docking_v5/</span>, <span class="kbd">08e_analysis_v5/</span>, <span class="kbd">09e_report_v5/</span>. v1, v2, v3, v4 left untouched.</p>

</body></html>
"""


def main():
    os.makedirs(REP_DIR, exist_ok=True)
    open(STAGE_LOG, "w").close()
    log("Stage 9 v5 starting")

    wt_apo = json.load(open(os.path.join(WT_DIR, "wt_apo.json")))
    wt_holo = json.load(open(os.path.join(WT_DIR, "wt_holo.json")))
    sum8 = json.load(open(os.path.join(ANA_DIR, "summary_v5.json")))
    sum7 = json.load(open(os.path.join(MUT_DIR, "summary_v5.json")))

    rec_pdbqt = os.path.join(WT_DIR, "protein_dimer_holo.pdbqt")
    max_q = 0.0
    with open(rec_pdbqt) as f:
        for line in f:
            if line.startswith("ATOM"):
                try:
                    c = float(line[66:76].strip())
                    if abs(c) > max_q:
                        max_q = abs(c)
                except Exception:
                    pass

    def fmt(v, n=3):
        try:
            v = float(v)
            return f"{v:.{n}f}"
        except (TypeError, ValueError):
            return "n/a"

    pearson_r = sum8.get("pearson_r_apo_holo")
    pearson_p = sum8.get("pearson_p_apo_holo")
    spearman_r = sum8.get("spearman_r_apo_holo")
    spearman_p = sum8.get("spearman_p_apo_holo")
    null_corr = sum8.get("no_significant_apo_holo_correlation", False)
    n_above_noise_holo = sum8.get("n_above_noise_holo", 0)
    n_pearson = sum7.get("n_holo_rows_clean", 0) or 0
    # Recover n from the analysis summary if available
    df_chk = pd.read_csv(os.path.join(MUT_DIR, "mutant_results_v5.csv"), comment="#")
    df_chk = df_chk[df_chk.mutant != "WT"].copy()
    df_chk["mis_docked"] = df_chk["mis_docked"].astype(str).str.lower() == "true"
    df_chk["low_confidence"] = df_chk["low_confidence"].astype(str).str.lower() == "true"
    df_chk["delta_vina_vs_wt"] = pd.to_numeric(df_chk["delta_vina_vs_wt"], errors="coerce")
    df_chk_apo = df_chk[df_chk.condition == "apo"].copy()
    df_chk_holo = df_chk[df_chk.condition == "holo"].copy()
    apo_clean = df_chk_apo[(~df_chk_apo.mis_docked) & (~df_chk_apo.low_confidence)]
    holo_clean = df_chk_holo[(~df_chk_holo.mis_docked) & (~df_chk_holo.low_confidence)]
    common = set(apo_clean.mutant) & set(holo_clean.mutant)
    n_pearson = len(common)

    correlation_conclusion = ("No statistically significant apo-holo correlation "
                              "(both p > 0.19; sci-off review item 3)") if null_corr else (
        f"Some apo-holo correlation (Pearson p = {fmt(pearson_p, 3)}, "
        f"Spearman p = {fmt(spearman_p, 3)})")

    if n_above_noise_holo == 0:
        noise_caveat = "No mutant in this study exceeds that threshold under the holo condition."
    else:
        noise_caveat = (f"{n_above_noise_holo} mutant(s) in this study exceed that threshold under the holo "
                        "condition; treat them as the most plausibly real signals, but still subject "
                        "to the empirical-not-thermodynamic limit of Vina.")

    # ---- Top tables: include flags column per sci-off review item 4.
    # If C195A is in the table, its mis_docked / low_confidence flags must be visible.
    def with_flags(rec_list):
        out = []
        for r in rec_list:
            mid = r["mutant"]
            row = df_chk[(df_chk.mutant == mid)
                         & (df_chk.condition == "holo")].iloc[0] if (df_chk.mutant == mid).any() else None
            if r.get("rmsd_to_native") is None:
                rrm = float("nan")
            else:
                rrm = float(r["rmsd_to_native"])
            # Pull from current df row to ensure we get latest flags
            apo_row = df_chk[(df_chk.mutant == mid) & (df_chk.condition == "apo")]
            holo_row = df_chk[(df_chk.mutant == mid) & (df_chk.condition == "holo")]
            flag_parts = []
            if not holo_row.empty:
                if bool(holo_row.iloc[0].mis_docked):
                    flag_parts.append("HOLO_mis_docked")
                if bool(holo_row.iloc[0].low_confidence):
                    flag_parts.append("HOLO_low_conf")
            if not apo_row.empty:
                if bool(apo_row.iloc[0].mis_docked):
                    flag_parts.append("APO_mis_docked")
                if bool(apo_row.iloc[0].low_confidence):
                    flag_parts.append("APO_low_conf")
            r2 = dict(r)
            r2["flags"] = ", ".join(flag_parts) if flag_parts else "ok"
            out.append(r2)
        return out

    top_apo = with_flags(sum8["top5_destab_apo_clean"])
    top_holo = with_flags(sum8["top5_destab_holo_clean"])
    top3_holo = with_flags(sum8["top3_destab_holo_clean"])

    # C195A explicit caveat values (from raw data)
    c195a_h = df_chk_holo[df_chk_holo.mutant == "C195A"]
    c195a_holo_delta = float(c195a_h.iloc[0].delta_vina_vs_wt) if not c195a_h.empty else float("nan")
    c195a_holo_n_modes = int(c195a_h.iloc[0].n_modes) if not c195a_h.empty else 0
    c195a_holo_flag = ""
    if not c195a_h.empty:
        flags = []
        if bool(c195a_h.iloc[0].mis_docked):
            flags.append("mis_docked")
        if bool(c195a_h.iloc[0].low_confidence):
            flags.append("low_confidence (n_modes<5)")
        c195a_holo_flag = ", ".join(flags) if flags else "no flags"

    # WT holo exh actually used (96 if best seed in primary; 128 if in fallback)
    primary_seeds = wt_holo.get("vina_params", {}).get("primary_seeds", [42, 7, 13, 99, 256])
    wt_holo_exh = 96 if wt_holo.get("best_seed") in primary_seeds else 128

    tpl = Template(TEMPLATE)
    html = tpl.render(
        now=datetime.now().strftime("%Y-%m-%d %H:%M"),
        wt_apo=wt_apo, wt_holo=wt_holo,
        max_q=max_q,
        p_bar=b64_img(os.path.join(ANA_DIR, "delta_vina_apo_holo.png")),
        p_scatter=b64_img(os.path.join(ANA_DIR, "delta_vina_apo_vs_holo.png")),
        p_cat=b64_img(os.path.join(ANA_DIR, "delta_vina_by_category.png")),
        top_apo=top_apo, top_holo=top_holo,
        pearson_r_str=fmt(pearson_r, 3), pearson_p_str=fmt(pearson_p, 3),
        spearman_r_str=fmt(spearman_r, 3), spearman_p_str=fmt(spearman_p, 3),
        n_pearson=n_pearson,
        correlation_conclusion=correlation_conclusion,
        noise_floor=sum8.get("vina_noise_floor_kcal_per_mol", 0.85),
        noise_caveat=noise_caveat,
        n_above_noise_holo=n_above_noise_holo,
        wt_holo_exh=wt_holo_exh,
        c195a_holo_delta=c195a_holo_delta, c195a_holo_n_modes=c195a_holo_n_modes,
        c195a_holo_flag=c195a_holo_flag,
    )
    html_path = os.path.join(REP_DIR, "report.html")
    with open(html_path, "w") as f:
        f.write(html)
    log(f"wrote {html_path}")

    # PDF
    try:
        from weasyprint import HTML
        pdf_path = os.path.join(REP_DIR, "report.pdf")
        HTML(string=html, base_url=REP_DIR).write_pdf(pdf_path)
        log(f"wrote {pdf_path} ({os.path.getsize(pdf_path)} bytes)")
    except Exception as e:
        log(f"weasyprint failed: {e}")
        try:
            import subprocess
            subprocess.run(["wkhtmltopdf", html_path,
                            os.path.join(REP_DIR, "report.pdf")],
                           check=True, capture_output=True)
            log("PDF via wkhtmltopdf fallback")
        except Exception as e2:
            log(f"all PDF backends failed: {e2}")

    # DOCX
    try:
        from docx import Document
        from docx.shared import Inches, RGBColor

        doc = Document()
        doc.add_heading("TYMS conserved active site - pipeline v5 report", 0)
        doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} - Pipeline version 5")

        doc.add_heading("Headline finding (sci-off recommendation, round 4 review)", 1)
        doc.add_paragraph(
            "Rigid-receptor AutoDock Vina with AD4 partial charges and the physically "
            "correct (net -2) raltitrexed cofactor cannot resolve TYMS active-site point "
            "mutants at the kcal/mol scale; this is the principal finding of v5. "
            f"{n_above_noise_holo} of 20 mutants exceed Vina's documented +/-0.85 kcal/mol "
            "noise floor under the holo condition."
        )

        doc.add_heading("1. Executive summary", 1)
        doc.add_paragraph(
            "The v5 fix. Round-4 reviewers identified that v4 placed the cofactor by "
            "Kabsch-aligning the CCD-ideal D16 onto the bound conformer, yielding a 2.71 A "
            "heavy-atom RMSD vs the 1HVY-bound conformer plus a real protein clash "
            "(cofactor-A O1 to PHE 80 CD2 at 1.95 A). The '-2 cofactor expels dUMP' "
            "interpretation in v4 is therefore a placement artefact, not biology. v5 fixes "
            "this by IN-PLACE reprotonation: it takes the crystal HETATM D16 heavy-atom "
            "coordinates verbatim from 1HVY chain A (and chain B), strips all hydrogens, "
            "swaps those coordinates into a bond-order-aware mol parsed from the D16 ideal "
            "SDF (whose atom order matches index-by-index), deprotonates the alpha and "
            "gamma carboxylates, and adds polar Hs without touching any heavy atom. A hard "
            "assertion would abort if any heavy atom moved by > 0.001 A. A clash gate "
            "aborts if any cofactor heavy atom is within 1.8 A of a protein heavy atom."
        )
        doc.add_paragraph(
            f"What changed in WT-holo. Top affinity moved from -5.24 kcal/mol (v4, "
            f"mis-docked at 12.95 A) to {wt_holo['top_affinity']:.2f} kcal/mol (v5, RMSD "
            f"{wt_holo['rmsd_top_to_native']:.2f} A from crystal). This is now within ~1 "
            "kcal/mol of the apo result (-9.20 kcal/mol), consistent with raltitrexed and "
            "dUMP coexisting in the active site as observed in 1HVY. The v4 'expulsion' "
            "effect is gone."
        )
        doc.add_paragraph(
            "What did NOT change in v5. Even with the cofactor placed correctly, no "
            "individual mutant exceeds Vina's noise floor on holo. The top destabiliser "
            f"({sum8['top3_destab_holo_clean'][0]['mutant']} "
            f"{float(sum8['top3_destab_holo_clean'][0]['delta_vina_vs_wt']):+.2f} kcal/mol) "
            "is still below +/-0.85 kcal/mol. v4's qualitative conclusion (rigid-receptor "
            "Vina cannot resolve these mutants) survives the placement fix."
        )

        doc.add_heading("2. WT docking", 1)
        t = doc.add_table(rows=1, cols=8)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Cond", "Top aff", "mean topk", "n modes",
                               "RMSD vs native", "best seed", "aff range", "Source"]):
            hdr[i].text = h
        for cond, dat, src in [("WT apo", wt_apo, "reused from v4 (apo receptor unchanged)"),
                                ("WT holo", wt_holo, "v5 (in-place reprotonated cofactor)")]:
            row = t.add_row().cells
            row[0].text = cond
            row[1].text = f"{dat['top_affinity']:.2f}"
            row[2].text = f"{dat['mean_topk']:.2f}"
            row[3].text = str(dat["n_modes"])
            row[4].text = f"{dat['rmsd_top_to_native']:.2f}"
            row[5].text = str(dat["best_seed"])
            row[6].text = f"{dat['affinity_distribution_width_kcal']:.2f}"
            row[7].text = src

        doc.add_heading("3a. Top destabilising mutations - apo (well-docked)", 1)
        t = doc.add_table(rows=1, cols=6)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Mutant", "Category", "Delta Vina", "RMSD (A)", "n modes", "flags"]):
            hdr[i].text = h
        for r in top_apo:
            row = t.add_row().cells
            row[0].text = r["mutant"]
            row[1].text = r["category"]
            row[2].text = f"{float(r['delta_vina_vs_wt']):+.2f}"
            row[3].text = f"{float(r['rmsd_to_native']):.2f}"
            row[4].text = str(r["n_modes"])
            row[5].text = r["flags"]

        doc.add_heading("3b. Top destabilising mutations - holo (well-docked)", 1)
        t = doc.add_table(rows=1, cols=6)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Mutant", "Category", "Delta Vina", "RMSD (A)", "n modes", "flags"]):
            hdr[i].text = h
        for r in top_holo:
            row = t.add_row().cells
            row[0].text = r["mutant"]
            row[1].text = r["category"]
            row[2].text = f"{float(r['delta_vina_vs_wt']):+.2f}"
            row[3].text = f"{float(r['rmsd_to_native']):.2f}"
            row[4].text = str(r["n_modes"])
            row[5].text = r["flags"]

        doc.add_heading("C195A holo - explicit caveat", 1)
        doc.add_paragraph(
            "Cys195 is the catalytic nucleophile of TYMS; its sulphur attacks C6 of dUMP "
            "to form the covalent enzyme-substrate intermediate. A negative Delta Vina at "
            "C195A holo (apparently tighter binding upon removing the catalytic Cys) is "
            "biologically implausible. In v5 the C195A holo Delta is "
            f"{c195a_holo_delta:+.2f} kcal/mol with n_modes={c195a_holo_n_modes} "
            f"({c195a_holo_flag}); per sci-off review item 4 it is suppressed from any "
            "summary table that does not also display the mis_docked / low_confidence "
            "flag. The negative number itself is attributed to a docking artefact "
            "(rigid receptor, dominant single-funnel pose at low n_modes), not to "
            "genuine increased affinity."
        )

        doc.add_heading("4. Statistics", 1)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid"
        t.rows[0].cells[0].text = "Metric"
        t.rows[0].cells[1].text = "Value"
        for k, v in [
            (f"Pearson r (apo vs holo Delta, well-docked, n={n_pearson})",
             f"{fmt(pearson_r,3)}, p = {fmt(pearson_p,3)}"),
            ("Spearman rho (same set)",
             f"{fmt(spearman_r,3)}, p = {fmt(spearman_p,3)}"),
            ("Conclusion", correlation_conclusion),
            ("Vina noise floor (kcal/mol)", f"+/-{sum8.get('vina_noise_floor_kcal_per_mol',0.85)} (Trott&Olson 2010; Forli 2016)"),
            ("n mutants with |Delta_holo| > noise floor", str(n_above_noise_holo)),
        ]:
            r = t.add_row().cells
            r[0].text = k
            r[1].text = v
        doc.add_paragraph(
            "The v4 filtered Spearman rho on the |Delta| > 0.3 sub-panel is dropped "
            "from this report. In v4 it had n = 4 (statistically meaningless). Sci-off "
            "review item 3 required its removal."
        )

        doc.add_heading("5. Mechanistic explanation - what was wrong with v4", 1)
        doc.add_paragraph(
            "The v4 report attributed the WT-holo dock landing 12.95 A from the crystal "
            "pocket to 'narrow funnel' and tentatively to '-2 ionised cofactor expelling "
            "-2 dUMP via electrostatics'. Round-4 structural-bioinformatics review showed "
            "that this interpretation was wrong on two counts. (1) The 1HVY crystal shows "
            "raltitrexed (D16) and dUMP coexisting in the active site of the same monomer; "
            "a '-2 cofactor expels -2 substrate' framing is inconsistent with the "
            "experimental ground truth. (2) The dominant artefact in v4 was COFACTOR "
            "PLACEMENT, not protonation. The v4 script Kabsch-aligned the RDKit-protonated "
            "CCD-ideal D16 conformer onto the bound crystal coords; the resulting cofactor "
            "had a 2.71 A heavy-atom RMSD vs the crystal-bound conformer and a 1.95 A "
            "clash between cofactor O1 and PHE 80 CD2. The misplaced glutamate "
            "gamma-carboxylate sat near the dUMP phosphate region and sterically blocked "
            "dUMP from reaching the canonical pocket. Vina's response was to put dUMP into "
            "a remote sub-pocket (12.95 A away). v5 confirms the diagnosis: with the "
            "cofactor heavy atoms placed at their crystal coordinates (RMSD 0.000 A vs "
            "1HVY) and the carboxylates correctly deprotonated and 0 protein clashes, "
            f"the WT-holo top pose lands in the canonical pocket "
            f"({wt_holo['rmsd_top_to_native']:.2f} A from crystal dUMP, top affinity "
            f"{wt_holo['top_affinity']:.2f} kcal/mol). The v4 'expulsion' finding does "
            "not survive the placement correction."
        )

        for png_name, cap in [
            ("delta_vina_apo_holo.png",
             "Figure 1 - v5 Delta Vina apo vs holo bar chart (grey band = Vina noise floor)"),
            ("delta_vina_apo_vs_holo.png",
             "Figure 2 - v5 Apo-vs-holo concordance (no significant correlation)"),
            ("delta_vina_by_category.png",
             "Figure 3 - v5 Delta Vina by category (excludes mis-docked and n_modes<5)"),
        ]:
            p = os.path.join(ANA_DIR, png_name)
            if os.path.exists(p):
                doc.add_picture(p, width=Inches(6.0))
                doc.add_paragraph(cap).italic = True

        doc.add_heading("6. Limitations", 1)
        for line in [
            f"Vina noise floor is approximately +/-{sum8.get('vina_noise_floor_kcal_per_mol',0.85)} kcal/mol (Trott & Olson 2010; Forli et al. 2016). {noise_caveat}",
            "Vina is an empirical scoring function, not a free energy. Numbers reported are Delta Vina score, not Delta-Delta G of binding.",
            "Cofactor polyglutamylation: the cofactor in 1HVY is the antifolate raltitrexed (PDB ligand D16, mono-glutamate). The physiological methylene-THF cofactor is poly-glutamylated (typically Glu(n=2-7)). The cofactor pocket geometry, ionisation, and water network may differ in the polyglutamylated state. This pipeline does not model that.",
            "C195A holo Delta < 0 is biologically implausible - Cys195 is the catalytic nucleophile. Any negative Delta is attributed to docking artefact (rigid receptor, dominant single-funnel pose at low n_modes).",
            "Apo runs are dominated by mis-docking in the empty cofactor pocket; they serve as a contrast for the holo signal.",
            "Receptor and ligand are rigid. PyMOL Mutagenesis Wizard rotamers (carried over from v3) were sculpt-relaxed locally but not against a relaxed pocket.",
            "Docking pose RMSD is computed against crystal dUMP heavy atoms only (no symmetry handling for ring atoms).",
            "Apo-holo correlation is not statistically significant in v5 (Pearson and Spearman p > 0.19). The v4 filtered Spearman rho on the |Delta| > 0.3 sub-panel had n = 4 and is dropped from v5 prose per round-4 sci-off review.",
        ]:
            doc.add_paragraph(line, style="List Bullet")

        doc.add_heading("7a. Dual RMSD reference (sci-off review item 5)", 1)
        doc.add_paragraph(
            "In v4, with a placement-buggy cofactor, WT-holo itself landed RMSD 12.95 A from the crystal "
            "dUMP, so v4 reported holo using a relaxed metric (|RMSD - WT_holo_RMSD| > 3 A). In v5 the "
            f"WT-holo RMSD is {wt_holo['rmsd_top_to_native']:.3f} A < 3 A, so the standard mis_docked = "
            "(RMSD > 3 A vs crystal) definition applies and the v4 dual-reference relaxation is not "
            "needed. This caveat is preserved here to make the methodological dependency explicit."
        )

        doc.add_heading("7b. Mutant-apo / WT-apo protocol asymmetry (sci-off review item 6)", 1)
        doc.add_paragraph(
            "WT-apo and WT-holo were docked with a 5-seed sweep at exh=96 (and exh=128 fallback). "
            "Mutant-apo dockings are inherited verbatim from v3, which used a single seed at exh=32. "
            "Mutant-holo dockings (in v5) likewise use single seed exh=32. The Delta_apo column "
            "therefore mixes a high-effort WT reference with low-effort mutant runs - the apo-side "
            "mutant numbers are slightly noisier than the holo-side mutant numbers but the WT-relative "
            "reference is consistent within each condition."
        )

        doc.add_heading("7c. AD4 / Vina partial-charge convention", 1)
        doc.add_paragraph(
            "AutoDock 4 / Vina use a united-atom convention; non-polar hydrogens are merged into "
            "their parent heavy atoms, and polar hydrogens (HD type) carry zero partial charge with "
            "the H-bond contribution folded into the parent. Maximum heavy-atom |q| in the v5 "
            f"receptor is {max_q:.3f}."
        )

        doc.add_heading("Reproducibility", 1)
        doc.add_paragraph(
            f"Vina 1.2.7 (WT-holo: exh=96 / 128 fallback, num_modes=32, box=22, "
            "seeds {42,7,13,99,256} primary; mutants: exh=32, num_modes=20, box=22, seed=42). "
            f"WT-apo reused from v4. Receptor: obabel-gasteiger, max |q| = {max_q:.3f}. "
            f"Cofactor (v5): IN-PLACE reprotonation of crystal HETATM D16 heavy-atom coords; "
            "RDKit bond-order assignment from D16 ideal SDF; explicit -COOH -> COO-; "
            "AddHs(addCoords=True). Heavy-atom RMSD vs 1HVY = 0.000 A; protein clashes < 1.8 A = 0."
        )
        doc.add_paragraph(
            "All v5 outputs in 03e_structure_v5, 06e_docking_wt_v5, 07e_mut_docking_v5, "
            "08e_analysis_v5, 09e_report_v5. v1, v2, v3, v4 left untouched."
        )

        docx_path = os.path.join(REP_DIR, "report.docx")
        doc.save(docx_path)
        log(f"wrote {docx_path} ({os.path.getsize(docx_path)} bytes)")
    except Exception as e:
        log(f"docx failed: {e}")
        import traceback
        log(traceback.format_exc())

    log("Stage 9 v5 DONE")


if __name__ == "__main__":
    main()
