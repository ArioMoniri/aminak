#!/usr/bin/env python3
"""
Build a self-contained Word (.docx) version of the pipeline report with all
images embedded as native Word pictures (not data URIs).

Run from project root, or pass project root as argv[1].
"""
from __future__ import annotations
import os, sys, csv, json, datetime, pathlib
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(sys.argv[1] if len(sys.argv) > 1 else os.environ.get("PROJECT_DIR", os.path.expanduser("~/conserved_site_project")))
assert ROOT.exists(), f"project root missing: {ROOT}"
OUT = ROOT / "09_report" / "report.docx"

doc = Document()

# Page setup
section = doc.sections[0]
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H(level: int, text: str):
    h = doc.add_heading(text, level=level)
    if level == 0:
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1f, 0x3b, 0x5e)
    return h


def P(text: str = "", *, italic=False, bold=False, size: int | None = None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p


def IMG(path: pathlib.Path, width_in: float = 6.5, caption: str | None = None):
    if not path.exists():
        P(f"[missing image: {path}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)


def TABLE(headers: list[str], rows: list[list[str]], col_widths_in: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            r = para.add_run(str(val))
            r.font.size = Pt(9)
    if col_widths_in:
        for row in t.rows:
            for i, w in enumerate(col_widths_in):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)


# ===================== TITLE =====================
H(0, "Conserved-Site Structural-Bioinformatics Pipeline — TYMS / dUMP Worked Example")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"Generated {datetime.date.today().isoformat()} · Human Thymidylate Synthase (UniProt P04818, PDB 1HVY) · Natural substrate dUMP").italic = True

# ===================== EXECUTIVE SUMMARY =====================
H(1, "Executive summary")
P(
    "End-to-end automated structural-bioinformatics pipeline that takes a single small/medium "
    "enzyme from a UniProt accession to a conservation-aware docking-and-mutagenesis report. "
    "The worked example is human Thymidylate Synthase (TYMS), the molecular target of "
    "5-fluorouracil in colorectal-cancer chemotherapy, docked against its natural substrate "
    "dUMP from PDB 1HVY (1.9 Å), with a 21-mutant probe panel and a distant-surface negative "
    "control. The pipeline executes 9 stages headlessly, produces 31 ray-traced PNGs, a full "
    "results table, and an HTML/PDF/DOCX report. The pipeline was subsequently audited by four "
    "specialised review agents (validator, code reviewer, scientific officer, structural "
    "bioinformatician); their verdicts and the residual limitations are reproduced verbatim "
    "in the dedicated section below."
)

# ===================== WORKFLOW DIAGRAM =====================
H(1, "Workflow")
P("Nine sequential stages, each isolated in its own numbered subfolder.")
IMG(ROOT / "workflow_diagram.png", width_in=6.5,
    caption="Pipeline workflow. Each stage is implemented as one Python script in scripts/.")

# ===================== TARGET CHOICE =====================
H(1, "1. Target & rationale")
P(
    "Human TYMS (UniProt P04818, 313 aa) is the molecular target of 5-fluorouracil (5-FU), "
    "the backbone of CRC chemotherapy for >50 years. PDB 1HVY (1.9 Å resolution) is the "
    "human enzyme co-crystallised with its natural substrate dUMP (residue name UMP) and "
    "the antifolate raltitrexed (residue D16) occupying the cofactor pocket. The dUMP-binding "
    "site is small (≤ 12 × 9 × 5 Å), well-defined, and its catalytic residues are "
    "experimentally characterised (Cys195 nucleophile, Arg phosphate clamp, His196/Asn226/Tyr135 "
    "proton/orientation network)."
)

# ===================== STAGE 1 — MSA =====================
H(1, "2. Cross-species multiple sequence alignment & conservation")
P(
    "Stage 1 fetches a panel of orthologs from UniProt by REST, aligns them with MAFFT "
    "(--auto), and computes a per-residue Jensen-Shannon divergence score (Capra & Singh "
    "2007 with Robinson-Robinson background; reference numbering = ungapped P04818)."
)
IMG(ROOT / "01_msa" / "conservation_plot.png", width_in=6.5,
    caption="Per-residue Jensen-Shannon conservation along human TYMS (top-10% peaks highlighted).")

# ===================== STAGE 2 — ACTIVE SITE =====================
H(1, "3. Active-site identification & overlap with conservation")
P(
    "Stage 2 collects active-site residues from two database sources — UniProt features "
    "(Active site / Binding site / Site) and the PDBe ligand-binding-graph API for 1HVY — "
    "and intersects them with the top-25 % of conserved residues from Stage 1."
)
IMG(ROOT / "02_active_site" / "overlap_figure.png", width_in=4.5,
    caption="Overlap between database-annotated active-site residues and conservation peaks.")

# ===================== STAGE 4 — VISUALIZATIONS =====================
H(1, "4. Structural visualisations")
P(
    "All four renders below were produced headlessly by the brewed PyMOL 3.1.0 binary via "
    "subprocess; no GUI was opened. Renders are 1600×1200 ray-traced."
)
for fn, cap in [
    ("01_overview.png", "Whole-protein surface + dUMP + active-site sticks (chain A)."),
    ("02_closeup.png", "Active-site closeup with residue labels."),
    ("03_conservation.png", "Cartoon coloured by Jensen-Shannon conservation (blue = low, red = high)."),
    ("04_cavity.png", "Surface cavity view (cavity_mode = 1)."),
]:
    IMG(ROOT / "04_pymol" / fn, width_in=5.5, caption=cap)

# ===================== STAGE 6 — WT DOCKING =====================
H(1, "5. Wild-type docking")
# Pull WT data
wt = json.loads((ROOT / "06_docking_wt" / "wt_result.json").read_text()) if (ROOT / "06_docking_wt" / "wt_result.json").exists() else {}
top_aff = wt.get("top_affinity")
mean3 = wt.get("mean_top3")
rmsd = wt.get("rmsd_to_native") or wt.get("rmsd_top_to_native")
P(
    f"AutoDock Vina 1.2.7. Box: 22³ Å centred on the active-site Cα centroid. Exhaustiveness 16, "
    f"num_modes 20, seed 42. WT top affinity: {top_aff} kcal/mol; mean of top-3: {mean3} kcal/mol; "
    f"top-pose heavy-atom RMSD vs crystal dUMP: {rmsd} Å (a credible positive control for the box)."
)
IMG(ROOT / "06_docking_wt" / "wt_topdock.png", width_in=6.0,
    caption="Wild-type receptor with the top docked dUMP pose overlaid on the crystal substrate.")

# ===================== STAGE 7 — MUTANTS =====================
H(1, "6. Mutational probe panel")
P(
    "21 mutants generated with PyMOL's Mutagenesis Wizard (rotamer pick, no pocket relaxation): "
    "8 alanine-scan singles of the selected active-site set, 7 chemically-opposite singles, "
    "5 mechanism-motivated doubles, and 1 distant-surface negative control (T170A, ~18 Å from "
    "the centroid). Every mutant was redocked with the IDENTICAL grid box, exhaustiveness, "
    "and seed as the WT run for fair Δ comparison."
)

# Read the results CSV
csv_path = ROOT / "07_mut_docking" / "results_full.csv"
rows = []
if csv_path.exists():
    with open(csv_path) as f:
        r = csv.DictReader(f)
        for row in r:
            rows.append(row)

# Sort by delta (descending = most disruptive first)
rows_sorted = sorted(rows, key=lambda x: -float(x.get("delta_vs_wt", 0) or 0))

# Compact table — top columns only
table_rows = []
for row in rows_sorted:
    mid = row["mutation_id"]
    typ = row["type"].replace("_", " ")
    aff = f"{float(row['top_affinity']):.2f}"
    delta = f"{float(row['delta_vs_wt']):+.2f}"
    rmsd = f"{float(row['rmsd_top_to_native']):.2f}"
    table_rows.append([mid, typ, aff, delta, rmsd])
TABLE(["Mutant", "Type", "Top affinity (kcal/mol)", "Δ vs WT", "RMSD vs crystal (Å)"],
      table_rows, col_widths_in=[1.4, 1.8, 1.4, 0.9, 1.2])

# ===================== STAGE 8 — ANALYSIS PLOTS + TEXT =====================
H(1, "7. Analysis")
for fn, cap in [
    ("delta_affinity_bar.png", "Δ Vina score per mutant relative to wild type, colour-coded by mutation type."),
    ("residue_substitution_heatmap.png", "Single-mutant Δ Vina score heatmap (residue × substitution)."),
    ("conservation_vs_effect.png", "Per-residue conservation (Jensen-Shannon) vs |Δ Vina score|. The flat scatter (r ≈ 0.03) reflects the MSA panel limitation discussed in §10."),
]:
    IMG(ROOT / "08_analysis" / fn, width_in=6.0, caption=cap)

# Pull analysis.md
am = ROOT / "08_analysis" / "analysis.md"
if am.exists():
    H(2, "Written interpretation (verbatim from 08_analysis/analysis.md)")
    text = am.read_text()
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            H(2, block.lstrip("# ").strip())
        elif block.startswith("## "):
            H(3, block.lstrip("# ").strip())
        elif block.startswith("- "):
            for line in block.split("\n"):
                if line.startswith("- "):
                    p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            P(block)

# ===================== STAGE 9 — METHODS / VERSIONS =====================
H(1, "8. Methods & software versions")
TABLE(
    ["Component", "Version", "Source"],
    [
        ["Python", "3.11.9", "pyenv"],
        ["AutoDock Vina", "1.2.7 (CLI via subprocess)", "brewsci/bio/autodock-vina"],
        ["PyMOL (headless)", "3.1.0", "homebrew/core/pymol"],
        ["MAFFT", "v7.526 (2024-04)", "homebrew/core/mafft"],
        ["OpenBabel", "3.1.0", "homebrew/core/open-babel"],
        ["Biopython", "1.87", "pypi"],
        ["RDKit", "(from venv freeze)", "pypi"],
        ["Meeko", "0.7.1", "pypi"],
        ["WeasyPrint, Jinja2, python-docx", "report stack", "pypi"],
    ],
    col_widths_in=[2.0, 2.2, 2.4]
)
P("Full library manifest in 00_setup/installed_libraries.md and 00_setup/pip_freeze.txt.")

# ===================== LIMITATIONS / REVIEWER FINDINGS =====================
H(1, "9. Multi-agent audit & known limitations")
P(
    "Four specialised review agents were spawned in parallel (read-only) after the pipeline "
    "completed, each grading a different facet of the output. Their full reports live in the "
    "reviews/ directory of the project; this section reproduces only the verdict line and the "
    "load-bearing critical findings."
)

TABLE(
    ["Reviewer", "Mandate", "Verdict"],
    [
        ["Validator", "File integrity, number reproduction, screenshot reality", "PASS with 2 ❌ + 5 ⚠️ flags"],
        ["Code Reviewer", "Python correctness, robustness, reproducibility", "12-item punch list (no blockers)"],
        ["Scientific Officer", "Peer-review-grade defensibility", "Needs major revision"],
        ["Structural Bioinformatician", "Deep technical methods audit", "FAIL"],
    ],
    col_widths_in=[2.4, 2.8, 1.8]
)

H(2, "Critical issues that materially affect interpretation")
for line in [
    "1. The MSA panel was built from the wrong proteins. Of the 9 UniProt accessions used in "
    "scripts/stage1_msa.py, only P04818 (human) and P07607 (mouse) are real TYMS. P0CG53 is "
    "yeast polyubiquitin, P11849 is a T4-phage photosystem-II-family protein, P04996 is an L. "
    "casei membrane protein, P04394 is E. coli ATP synthase. The Jensen-Shannon scores in "
    "01_msa/conservation_scores.csv are therefore noise, and the Stage-2 'top-25%-conserved' "
    "intersection had to be augmented manually with literature-known catalytic Cys195/His196 — "
    "a step which would not be needed against a real TYMS alignment.",
    "2. TYMS is an obligate homodimer; chain B was discarded. Catalytic residues from the "
    "partner subunit (Arg175 / Arg176 clamp the dUMP phosphate) are missing from the docking "
    "receptor. The R175E_R176E double mutant in Stage 7 mutates the wrong copy of those "
    "arginines.",
    "3. One mutant (G217W) has 9 atomic clashes, the worst at 0.98 Å between Trp CD1 and "
    "Val223 CG2. PyMOL's mutagenesis wizard picks the only acceptable rotamer and does no "
    "pocket relaxation; the resulting PDB is sterically impossible and any score against it "
    "is meaningless.",
    "4. CME43 (covalently modified Cys) was silently dropped during structure cleanup, "
    "leaving a 42→44 backbone gap.",
    "5. Cys195 is included in the active-site set despite a JSD percentile of only 36.7 — a "
    "consequence of issue #1, not a feature of the enzyme.",
]:
    p = doc.add_paragraph(line, style="List Number")

H(2, "What still stands")
P(
    "The docking + mutagenesis half (rigid-receptor Vina, single chain, Gasteiger charges, "
    "no minimisation) is internally self-consistent: WT redock RMSD is 1.08 Å, the distant-"
    "surface negative control T170A gives ΔVina ≈ 0, and the pose-displacing mutants D218K "
    "(RMSD 7.34 Å), Y258A (4.51 Å), and N226A (5.88 Å) behave the way a competent rigid-"
    "receptor pipeline would predict. The C195 'scores better than WT' result is acknowledged "
    "in §7 as a rigid-receptor artefact and not over-interpreted."
)

H(2, "Required revisions before external use")
for line in [
    "Replace 01_msa/input.fa with verified TYMS orthologs only (correct accessions: P00469 "
    "L. casei, P0A884 E. coli, P07807 S. cerevisiae CDC21, Q23381 C. elegans, P13922 P. "
    "falciparum trimmed to TS domain, P04019 T4 td, plus Arabidopsis Q05762 and an archaeon). "
    "Re-run Stages 1, 2, and 8.",
    "Keep both chains A and B in scripts/stage3_structure.py. Re-run Stages 3–7.",
    "Drop or rebuild G217W with side-chain relaxation (PyRosetta or FoldX RepairPDB).",
    "Preserve CME43 by re-mutating to CYS (rename, strip 2-hydroxyethyl atoms) instead of "
    "dropping the residue.",
    "Switch ligand charges to AM1-BCC (Meeko mk_prepare_ligand) as the primary path.",
    "Add per-mutant prep-status row in the report so silent meeko/obabel fallbacks are "
    "visible.",
    "Add R175A, R176A, R215A, R50A single Ala-scans so the Arg phosphate-clamp claim is "
    "actually probed.",
]:
    p = doc.add_paragraph(line, style="List Bullet")

# ===================== FOOTER =====================
H(1, "10. Files in this report")
P("Master log: pipeline.log")
P("Master numerical table: 07_mut_docking/results_full.csv")
P("Reviewer reports: reviews/01_validator.md, reviews/02_code_review.md, reviews/03_scientific_officer.md, reviews/04_structural_bioinformatician.md")
P("Library manifest: 00_setup/installed_libraries.md")
P("Source: scripts/stage1_msa.py … stage9b_docx.py")

doc.save(str(OUT))
print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")
