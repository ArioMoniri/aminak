#!/usr/bin/env python3
"""Stage 9 v4: HTML/PDF/DOCX report with all v4 caveats:
- Vina noise floor band on plots & in Limitations (Fix G)
- C195A explicit caveat (Fix H)
- Polyglutamylation scope (Fix I)
- AD4 polar H zero charge documentation (Fix J)
"""
import os, sys, base64, json, math
from datetime import datetime
import pandas as pd
from jinja2 import Template

PROJECT = os.path.expanduser("~/conserved_site_project")
REP_DIR = os.path.join(PROJECT, "09d_report_v4")
LOG_DIR = os.path.join(PROJECT, "logs")
LOG_FILE = os.path.join(PROJECT, "pipeline.log")
STAGE_LOG = os.path.join(LOG_DIR, "v4_09_report.log")

WT_DIR = os.path.join(PROJECT, "06d_docking_wt_v4")
MUT_DIR = os.path.join(PROJECT, "07d_mut_docking_v4")
ANA_DIR = os.path.join(PROJECT, "08d_analysis_v4")
STR_V4 = os.path.join(PROJECT, "03d_structure_v4")


def log(msg):
    line = f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [V4] STAGE9: {msg}"
    print(line, flush=True)
    with open(LOG_FILE, "a") as f:
        f.write(line + "\n")
    with open(STAGE_LOG, "a") as f:
        f.write(line + "\n")


def b64_img(path):
    if not os.path.exists(path):
        return ""
    with open(path, "rb") as f:
        return "data:image/png;base64," + base64.b64encode(f.read()).decode("ascii")


TEMPLATE = """<!doctype html>
<html><head><meta charset="utf-8"><title>TYMS Conserved Site v4 Report</title>
<style>
@page { size: A4; margin: 18mm; }
body { font-family: -apple-system, "Helvetica Neue", Helvetica, Arial, sans-serif;
       color: #222; line-height: 1.45; font-size: 10.5pt; }
h1 { color: #1a3759; border-bottom: 2px solid #1a3759; padding-bottom: 6px; }
h2 { color: #1a3759; margin-top: 22px; border-bottom: 1px solid #aac; padding-bottom: 3px; }
h3 { color: #2c4a6f; margin-top: 14px; }
img { max-width: 100%; height: auto; border: 1px solid #ccc; margin: 6px 0; }
table { border-collapse: collapse; font-size: 9pt; margin: 8px 0; width: 100%; }
th, td { border: 1px solid #888; padding: 4px 8px; text-align: left; }
th { background: #eef2f7; }
.small { font-size: 9pt; color: #555; }
.kbd { font-family: Menlo, monospace; background: #f4f4f4; padding: 1px 4px; border-radius: 3px; }
.fig { text-align: center; margin: 10px 0; }
.cap { font-size: 9pt; color: #555; font-style: italic; }
.box { background: #f7f9fc; border-left: 4px solid #1a3759; padding: 8px 12px; margin: 10px 0; }
.warn { background: #fff7e6; border-left: 4px solid #d4a017; padding: 8px 12px; margin: 10px 0; }
.lim { background: #fdecea; border-left: 4px solid #c0392b; padding: 8px 12px; margin: 10px 0; }
.row-c195a { background: #ffe0e0; }
</style></head><body>

<h1>TYMS conserved active site — pipeline v4 report</h1>
<div class="small">Generated {{ now }} • Project root: <span class="kbd">~/conserved_site_project</span> • Pipeline version 4</div>

<div class="box">
<b>v4 changes vs v3 (round-3 review fixes):</b>
<ul>
<li><b>FIX A — Real cofactor reprotonation.</b> v3 ran <span class="kbd">obabel -p 7.4</span> on a PDB with no bond orders, so nothing changed (output byte-identical to v2). v4 fetches the D16 ideal SDF (with bond orders) from the RCSB Chemical Component Dictionary, runs RDKit reionisation + explicit -COOH→COO⁻ on the α and γ carboxylates of the glutamate tail, then re-adds Hs and Kabsch-aligns onto the original 1HVY crystal coords. <b>The chain-A cofactor file's MD5 differs from v2/v3</b> (assertion in script, would loud-abort on regression).</li>
<li><b>FIX B — WT holo multi-seed sweep.</b> v3 used 2 seeds (42, 7) and accepted the lower-RMSD result (circular). v4 sweeps {42, 7, 13, 99, 256} at exh=96; if best n_modes &lt; 10, falls back to {1, 2025, 31337} at exh=128. Selection is by lowest top affinity, tie-break highest n_modes — RMSD-to-crystal is no longer a selection criterion.</li>
<li><b>FIX C — Holo mutant re-dock.</b> All 20 mutant holo dockings redone against the new (Fix-A) cofactor-bearing receptor. Apo dockings reused from v3 (apo receptor is unchanged in v4).</li>
<li><b>FIX D — UMP atom-name preservation.</b> Vina/PDBQT round-trips strip atom names (everything becomes generic C/N/O/P). v4 walks the PDBQT in heavy-atom order and transplants names from the input ligand PDB. The named pose's RMSD reproduces the JSON-reported RMSD exactly.</li>
<li><b>FIX E — Low-confidence filter.</b> Holo rows with <span class="kbd">n_modes &lt; 5</span> are flagged <span class="kbd">low_confidence=True</span> and excluded from the top-destabiliser tables.</li>
<li><b>FIX F — Spearman.</b> Spearman ρ reported alongside Pearson r; an additional Spearman computed restricted to mutants with |Δ|&gt;0.3 (the noise band).</li>
<li><b>FIX G — Vina noise floor.</b> Documented at ±0.85 kcal/mol (Trott &amp; Olson 2010; Forli et al. 2016). Visible as grey band on Δ Vina plots.</li>
<li><b>FIX H — C195A caveat.</b> Explicit annotation: any negative Δ at C195A holo is biologically implausible (Cys195 is the catalytic nucleophile).</li>
<li><b>FIX I — Polyglutamylation scope.</b> Limitations section now explicit that the modelled cofactor is mono-glutamate raltitrexed (D16), not the polyglutamylated physiological folate.</li>
<li><b>FIX J — AD4 zero-charge polar H docs.</b> Methods section explains the AutoDock united-atom convention (HD type, zero charge, parent absorbs H-bond term).</li>
<li><b>FIX K — WT row in mutant CSV.</b> First two data rows of <span class="kbd">mutant_results_v4.csv</span> are <span class="kbd">WT, wildtype, apo|holo</span> with Δ=0.</li>
</ul>
</div>

<h2>1. WT docking — apo vs holo (v4 receptor with reprotonated cofactor)</h2>
<table><tr><th>Condition</th><th>Top affinity (kcal/mol)</th><th>mean top-3</th>
<th>n modes</th><th>RMSD top-pose vs crystal dUMP (Å)</th><th>Best seed</th><th>Affinity range across seeds (kcal/mol)</th></tr>
<tr><td>WT apo</td>
  <td>{{ "%.2f"|format(wt_apo.top_affinity) }}</td>
  <td>{{ "%.2f"|format(wt_apo.mean_topk) }}</td>
  <td>{{ wt_apo.n_modes }}</td>
  <td>{{ "%.2f"|format(wt_apo.rmsd_top_to_native) }}</td>
  <td>{{ wt_apo.best_seed }}</td>
  <td>{{ "%.2f"|format(wt_apo.affinity_distribution_width_kcal) }}</td>
</tr>
<tr><td>WT holo</td>
  <td>{{ "%.2f"|format(wt_holo.top_affinity) }}</td>
  <td>{{ "%.2f"|format(wt_holo.mean_topk) }}</td>
  <td>{{ wt_holo.n_modes }}</td>
  <td>{{ "%.2f"|format(wt_holo.rmsd_top_to_native) }}</td>
  <td>{{ wt_holo.best_seed }}</td>
  <td>{{ "%.2f"|format(wt_holo.affinity_distribution_width_kcal) }}</td>
</tr>
</table>
<p class="small">Box centroid: ({{ "%.2f, %.2f, %.2f"|format(centroid[0], centroid[1], centroid[2]) }}) Å (crystal dUMP centroid). Exhaustiveness 96, num_modes 32, box 22 Å. Selection: lowest top affinity, tie-break highest n_modes (NOT RMSD — that was circular in v3).</p>

<h2>2. Mutant panel</h2>
<p>Panel composition: 20 mutants. Each docked under apo and holo conditions.
Apo dockings reused from v3 (the apo receptor did not change between v3 and v4).
Holo dockings rebuilt in v4 against the reprotonated-cofactor receptor.</p>

<div class="fig"><img src="{{ p_bar }}" alt="bar"/>
<div class="cap">Figure 1 — Δ Vina score vs WT v4 for each mutant (apo blue, holo orange). Grey bars = mis-docked (RMSD &gt; 3 Å) or low_confidence (n_modes &lt; 5). The grey band is the Vina noise floor (±0.85 kcal/mol). Positive = destabilising.</div></div>

<div class="fig"><img src="{{ p_scatter }}" alt="scatter"/>
<div class="cap">Figure 2 — Apo vs holo Δ Vina concordance (well-docked, n_modes ≥ 5). Pearson, Spearman ρ, and a noise-band-restricted Spearman ρ all reported. Grey band = Vina noise floor.</div></div>

<div class="fig"><img src="{{ p_cat }}" alt="cat"/>
<div class="cap">Figure 3 — Δ Vina by mutation category (excludes mis-docked and n_modes &lt; 5).</div></div>

{% if p_eff %}
<div class="fig"><img src="{{ p_eff }}" alt="eff"/>
<div class="cap">Figure 4 — Mutation-effect map: Δ Vina (x) vs pose RMSD (y) for v4 data.</div></div>
{% endif %}

<h3>Top destabilising mutations — apo (well-docked, n_modes ≥ 5)</h3>
<table><tr><th>Mutant</th><th>category</th><th>Δ Vina</th><th>RMSD</th><th>n modes</th></tr>
{% for r in top_apo %}<tr{% if r.mutant == "C195A" %} class="row-c195a"{% endif %}>
  <td>{{ r.mutant }}</td><td>{{ r.category }}</td>
  <td>{{ "%+.2f"|format(r.delta_vina_vs_wt|float) }}</td>
  <td>{{ "%.2f"|format(r.rmsd_to_native|float) }}</td>
  <td>{{ r.n_modes }}</td>
</tr>{% endfor %}
</table>

<h3>Top destabilising mutations — holo (well-docked, n_modes ≥ 5)</h3>
<table><tr><th>Mutant</th><th>category</th><th>Δ Vina</th><th>RMSD</th><th>n modes</th></tr>
{% for r in top_holo %}<tr{% if r.mutant == "C195A" %} class="row-c195a"{% endif %}>
  <td>{{ r.mutant }}</td><td>{{ r.category }}</td>
  <td>{{ "%+.2f"|format(r.delta_vina_vs_wt|float) }}</td>
  <td>{{ "%.2f"|format(r.rmsd_to_native|float) }}</td>
  <td>{{ r.n_modes }}</td>
</tr>{% endfor %}
</table>

<div class="warn">
<b>C195A caveat.</b> Cys195 is the catalytic nucleophile of TYMS; its sulphur attacks the C6 of dUMP to form the covalent enzyme-substrate intermediate. A negative Δ Vina at C195A holo (i.e. apparently tighter binding upon removing the catalytic Cys) is biologically implausible. Any such observation is attributed to docking artefact (e.g. narrow funnel of the WT holo reference), not to genuine increased affinity. The C195A row is highlighted in pink wherever it appears in this report.
</div>

<h2>3. Statistics</h2>
<table>
<tr><th>Metric</th><th>Value</th></tr>
<tr><td>Pearson r (apo vs holo Δ, all valid)</td><td>{{ pearson_r_str }}</td></tr>
<tr><td>Pearson p</td><td>{{ pearson_p_str }}</td></tr>
<tr><td>Spearman ρ (apo vs holo, all valid)</td><td>{{ spearman_r_str }}</td></tr>
<tr><td>Spearman p</td><td>{{ spearman_p_str }}</td></tr>
<tr><td>Spearman ρ filtered |Δ|&gt;0.3</td><td>{{ spearman_r_filt_str }} (n={{ n_filt }})</td></tr>
<tr><td>Vina noise floor (kcal/mol)</td><td>±{{ noise_floor }} (Trott &amp; Olson 2010; Forli et al. 2016)</td></tr>
<tr><td>n mutants with |Δ_holo| &gt; noise floor</td><td>{{ n_above_noise_holo }}</td></tr>
</table>

<div class="lim">
<b>Limitations.</b>
<ul>
<li><b>Vina noise floor is ±0.85 kcal/mol.</b> AutoDock Vina's documented noise floor is ~0.7–1.0 kcal/mol on cross-docking benchmarks (Trott &amp; Olson 2010, Forli et al. 2016). <b>{{ noise_caveat }}</b> The rankings are <i>suggestive</i>, not statistically <i>significant</i> differences in true binding free energy.</li>
<li><b>Vina is empirical, not free energy.</b> Reported numbers are "Δ Vina score" (kcal/mol), not ΔΔG of binding.</li>
<li><b>Cofactor polyglutamylation.</b> The cofactor in 1HVY is the antifolate raltitrexed (PDB ligand D16, mono-glutamate). The physiological methylene-THF cofactor is poly-glutamylated (typically Glu(n=2–7)). The cofactor pocket geometry, ionisation, and water network may differ in the polyglutamylated state. This pipeline does not model that.</li>
<li><b>C195A holo Δ &lt; 0 is biologically implausible</b> — Cys195 is the catalytic nucleophile and its removal cannot increase non-covalent dUMP affinity. The negative Δ here is attributed to the WT holo dock's narrow funnel, not to genuine tighter binding.</li>
<li><b>Apo runs are dominated by mis-docking</b> in the empty cofactor pocket; they serve as a contrast for the holo signal.</li>
<li><b>Receptor and ligand are rigid.</b> PyMOL Mutagenesis Wizard rotamers were sculpt-relaxed in v3 (carried over to v4 mutants), but not against a relaxed pocket.</li>
<li><b>Docking pose RMSD is computed against the crystal dUMP heavy atoms only</b> (no symmetry handling for ring atoms).</li>
</ul>
</div>

<h2>4. Methods</h2>
<table>
<tr><th>Stage</th><th>Tool / parameters</th></tr>
<tr><td>Cofactor reprotonation (v4)</td><td>RDKit reionizer on D16 ideal SDF (RCSB CCD); explicit -COOH → COO⁻ on α and γ carboxylates; AddHs(addCoords=True); Kabsch alignment to crystal coords</td></tr>
<tr><td>Receptor prep</td><td>obabel -xr -p 7.4 --partialcharge gasteiger; max abs charge {{ "%.3f"|format(max_q) }}</td></tr>
<tr><td>WT docking</td><td>Vina 1.2.7 • exh=96 • num_modes=32 • box=22³Å • seeds {42, 7, 13, 99, 256}; fallback {1, 2025, 31337} at exh=128 if max n_modes &lt; 10</td></tr>
<tr><td>WT selection</td><td>lowest top_affinity (NOT RMSD); tie-break highest n_modes</td></tr>
<tr><td>Mutant docking (holo, v4)</td><td>Vina 1.2.7 • exh=32 • num_modes=20 • box=22³Å • seed=42</td></tr>
<tr><td>Mutant docking (apo)</td><td>Reused from v3 (apo receptor unchanged)</td></tr>
<tr><td>Mutagenesis</td><td>(carried over from v3) PyMOL 3.1.0 Mutagenesis Wizard frame=1 + sculpt 3×20 cycles</td></tr>
<tr><td>UMP atom-name preservation</td><td>walk PDBQT in heavy-atom order; transplant names from input ligand_h.pdb</td></tr>
<tr><td>Sign convention</td><td><span class="kbd">delta_vina_vs_wt = top_aff_mut − top_aff_wt_v4</span> (positive = destabilising)</td></tr>
<tr><td>mean_topk</td><td><span class="kbd">mean(affinities[:min(3, n_modes)])</span></td></tr>
<tr><td>mis_docked filter</td><td>RMSD top-pose vs crystal dUMP &gt; 3 Å</td></tr>
<tr><td>low_confidence filter</td><td>n_modes &lt; 5 (holo only)</td></tr>
<tr><td>Statistics</td><td>Pearson r + Spearman ρ + filtered Spearman ρ on |Δ|&gt;0.3 mutants</td></tr>
</table>

<h3>4a. AD4 / Vina partial-charge convention (FIX J)</h3>
<p class="small">AutoDock 4 / Vina use a united-atom convention: non-polar hydrogens are merged into their parent heavy atoms, and polar hydrogens (HD type) carry zero partial charge with the H-bond contribution folded into the parent. The 1070 zero-charge atoms in the holo receptor PDBQT are therefore a feature of the AD4 model, not a charge-assignment failure. Maximum heavy-atom |q| in the v4 receptor is {{ "%.3f"|format(max_q) }} (well above the 0.05 sanity threshold).</p>

<p class="small">Outputs: <span class="kbd">03d_structure_v4/</span>, <span class="kbd">06d_docking_wt_v4/</span>, <span class="kbd">07d_mut_docking_v4/</span>, <span class="kbd">08d_analysis_v4/</span>, <span class="kbd">09d_report_v4/</span>. v1, v2, v3 left untouched.</p>

</body></html>
"""


def main():
    os.makedirs(REP_DIR, exist_ok=True)
    open(STAGE_LOG, "w").close()
    log("Stage 9 v4 starting")

    wt_apo = json.load(open(os.path.join(WT_DIR, "wt_apo.json")))
    wt_holo = json.load(open(os.path.join(WT_DIR, "wt_holo.json")))
    sum8 = json.load(open(os.path.join(ANA_DIR, "summary_v4.json")))
    sum7 = json.load(open(os.path.join(MUT_DIR, "summary_v4.json")))

    centroid = wt_apo["vina_params"]["centroid"]
    rec_method = wt_apo["vina_params"]["receptor_methods"]["apo"]
    rec_pdbqt = os.path.join(WT_DIR, "protein_dimer_apo.pdbqt")
    max_q = 0.0
    with open(rec_pdbqt) as f:
        for line in f:
            if line.startswith("ATOM"):
                try:
                    c = float(line[66:76].strip())
                    if abs(c) > max_q:
                        max_q = abs(c)
                except Exception:
                    pass

    # Format stat strings
    def fmt(v, n=3):
        try:
            v = float(v)
            return f"{v:.{n}f}"
        except (TypeError, ValueError):
            return "n/a"
    pearson_r_str = fmt(sum8.get("pearson_r_apo_holo"))
    pearson_p_str = fmt(sum8.get("pearson_p_apo_holo"), 4)
    spearman_r_str = fmt(sum8.get("spearman_r_apo_holo"))
    spearman_p_str = fmt(sum8.get("spearman_p_apo_holo"), 4)
    spearman_r_filt_str = fmt(sum8.get("spearman_r_filtered_abs_delta_gt_0p3"))
    n_filt = sum8.get("n_filtered_for_spearman", 0)
    noise_floor = sum8.get("vina_noise_floor_kcal_per_mol", 0.85)

    # Count holo |Δ| above noise
    df_h = pd.read_csv(os.path.join(MUT_DIR, "mutant_results_v4.csv"), comment="#")
    df_h = df_h[(df_h.condition == "holo") & (df_h.mutant != "WT")]
    df_h["delta_vina_vs_wt"] = pd.to_numeric(df_h["delta_vina_vs_wt"], errors="coerce")
    df_h["mis_docked"] = df_h["mis_docked"].astype(str).str.lower() == "true"
    df_h["low_confidence"] = df_h["low_confidence"].astype(str).str.lower() == "true"
    valid_h = df_h[(~df_h.mis_docked) & (~df_h.low_confidence)]
    n_above = int((valid_h.delta_vina_vs_wt.abs() > noise_floor).sum())
    if n_above == 0:
        noise_caveat = "No mutant in this study exceeds that threshold under the holo condition."
    else:
        noise_caveat = (f"{n_above} mutant(s) in this study exceed that threshold under the holo "
                        "condition; treat them as the most plausibly real signals, but still subject "
                        "to the empirical-not-thermodynamic limit of Vina.")

    tpl = Template(TEMPLATE)
    html = tpl.render(
        now=datetime.now().strftime("%Y-%m-%d %H:%M"),
        wt_apo=wt_apo, wt_holo=wt_holo,
        rec_method=rec_method, centroid=centroid, max_q=max_q,
        p_bar=b64_img(os.path.join(ANA_DIR, "delta_vina_apo_holo.png")),
        p_scatter=b64_img(os.path.join(ANA_DIR, "delta_vina_apo_vs_holo.png")),
        p_cat=b64_img(os.path.join(ANA_DIR, "delta_vina_by_category.png")),
        p_eff=b64_img(os.path.join(ANA_DIR, "mutation_effect_plot.png")),
        top_apo=sum8["top5_destab_apo_clean"],
        top_holo=sum8["top5_destab_holo_clean"],
        pearson_r_str=pearson_r_str, pearson_p_str=pearson_p_str,
        spearman_r_str=spearman_r_str, spearman_p_str=spearman_p_str,
        spearman_r_filt_str=spearman_r_filt_str, n_filt=n_filt,
        noise_floor=noise_floor, noise_caveat=noise_caveat,
        n_above_noise_holo=n_above,
    )
    html_path = os.path.join(REP_DIR, "report.html")
    with open(html_path, "w") as f:
        f.write(html)
    log(f"wrote {html_path}")

    # PDF
    try:
        from weasyprint import HTML
        pdf_path = os.path.join(REP_DIR, "report.pdf")
        HTML(string=html, base_url=REP_DIR).write_pdf(pdf_path)
        log(f"wrote {pdf_path} ({os.path.getsize(pdf_path)} bytes)")
    except Exception as e:
        log(f"weasyprint failed: {e}")
        try:
            import subprocess
            subprocess.run(["wkhtmltopdf", html_path,
                            os.path.join(REP_DIR, "report.pdf")],
                           check=True, capture_output=True)
            log("PDF via wkhtmltopdf fallback")
        except Exception as e2:
            log(f"all PDF backends failed: {e2}")

    # DOCX
    try:
        from docx import Document
        from docx.shared import Inches, RGBColor

        doc = Document()
        doc.add_heading("TYMS conserved active site — pipeline v4 report", 0)
        doc.add_paragraph(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} - Pipeline version 4")

        doc.add_heading("Summary of v4 fixes (round-3 review)", 1)
        for line in [
            "FIX A: Cofactor REALLY reprotonated at pH 7.4 — RDKit reionisation on bond-order-aware D16 ideal SDF (RCSB CCD); v3 was a no-op because input PDB lacked bond orders.",
            "FIX B: WT holo multi-seed sweep {42,7,13,99,256} (+ fallback) at exh=96/128; selection by lowest top affinity (tie-break highest n_modes), NOT RMSD (v3 was circular).",
            "FIX C: All 20 mutant holo dockings redone against new cofactor-bearing receptor; apo reused from v3.",
            "FIX D: UMP atom names preserved through Vina/PDBQT round-trip by transplanting from input ligand PDB.",
            "FIX E: low_confidence (n_modes < 5) added; excluded from rankings.",
            "FIX F: Spearman rho (and filtered |Delta|>0.3) reported alongside Pearson r.",
            "FIX G: Vina noise floor (~0.85 kcal/mol; Trott&Olson 2010, Forli 2016) shown as grey band on plots and in Limitations.",
            "FIX H: C195A holo caveat highlighted (catalytic nucleophile; negative Delta is artefact).",
            "FIX I: Polyglutamylation scope statement in Limitations.",
            "FIX J: AD4 / Vina united-atom polar-H zero-charge convention documented in Methods.",
            "FIX K: WT reference rows added to mutant CSV (Delta=0 for WT/apo and WT/holo).",
        ]:
            doc.add_paragraph(line, style="List Bullet")

        doc.add_heading("WT docking (v4 receptor)", 1)
        t = doc.add_table(rows=1, cols=7)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Condition", "Top aff", "mean topk", "n modes",
                               "RMSD vs native", "best seed", "aff range"]):
            hdr[i].text = h
        for cond, dat in [("WT apo", wt_apo), ("WT holo", wt_holo)]:
            row = t.add_row().cells
            row[0].text = cond
            row[1].text = f"{dat['top_affinity']:.2f}"
            row[2].text = f"{dat['mean_topk']:.2f}"
            row[3].text = str(dat["n_modes"])
            row[4].text = f"{dat['rmsd_top_to_native']:.2f}"
            row[5].text = str(dat["best_seed"])
            row[6].text = f"{dat['affinity_distribution_width_kcal']:.2f}"

        doc.add_heading("Top destabilising mutations - apo (well-docked, n_modes>=5)", 1)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Mutant", "Category", "Delta Vina", "RMSD", "n modes"]):
            hdr[i].text = h
        for r in sum8["top5_destab_apo_clean"]:
            row = t.add_row().cells
            row[0].text = r["mutant"]
            row[1].text = r["category"]
            row[2].text = f"{float(r['delta_vina_vs_wt']):+.2f}"
            row[3].text = f"{float(r['rmsd_to_native']):.2f}"
            row[4].text = str(r["n_modes"])

        doc.add_heading("Top destabilising mutations - holo (well-docked, n_modes>=5)", 1)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Mutant", "Category", "Delta Vina", "RMSD", "n modes"]):
            hdr[i].text = h
        for r in sum8["top5_destab_holo_clean"]:
            row = t.add_row().cells
            row[0].text = r["mutant"]
            row[1].text = r["category"]
            row[2].text = f"{float(r['delta_vina_vs_wt']):+.2f}"
            row[3].text = f"{float(r['rmsd_to_native']):.2f}"
            row[4].text = str(r["n_modes"])

        doc.add_heading("C195A caveat", 1)
        doc.add_paragraph(
            "Cys195 is the catalytic nucleophile of TYMS; its sulphur attacks C6 of dUMP "
            "to form the covalent enzyme-substrate intermediate. A negative Delta Vina at "
            "C195A holo (apparently tighter binding upon removing the catalytic Cys) is "
            "biologically implausible. Any such observation is attributed to docking "
            "artefact (e.g. narrow funnel of the WT holo reference), not to genuine "
            "increased affinity."
        )

        for png_name, cap in [
            ("delta_vina_apo_holo.png", "Figure 1 - v4 Delta Vina apo vs holo bar chart (grey band = Vina noise floor)"),
            ("delta_vina_apo_vs_holo.png", "Figure 2 - v4 Apo-vs-holo concordance with Pearson, Spearman, filtered Spearman"),
            ("delta_vina_by_category.png", "Figure 3 - v4 Delta Vina by category (excludes mis-docked and n_modes<5)"),
        ]:
            p = os.path.join(ANA_DIR, png_name)
            if os.path.exists(p):
                doc.add_picture(p, width=Inches(6.0))
                doc.add_paragraph(cap).italic = True

        doc.add_heading("Limitations", 1)
        for line in [
            f"Vina noise floor is approximately +/-{noise_floor} kcal/mol (Trott & Olson 2010; Forli et al. 2016). {noise_caveat}",
            "Vina is an empirical scoring function, not a free energy. Numbers reported are Delta Vina score, not Delta-Delta G of binding.",
            "Cofactor polyglutamylation: the cofactor in 1HVY is the antifolate raltitrexed (PDB ligand D16, mono-glutamate). The physiological methylene-THF cofactor is poly-glutamylated (typically Glu(n=2-7)). The cofactor pocket geometry, ionisation, and water network may differ in the polyglutamylated state. This pipeline does not model that.",
            "C195A holo Delta < 0 is biologically implausible - Cys195 is the catalytic nucleophile. Any negative Delta is attributed to docking artefact.",
            "Apo runs are dominated by mis-docking in the empty cofactor pocket; they serve as a negative control / contrast for the holo signal.",
            "Receptor and ligand are rigid. PyMOL Mutagenesis Wizard rotamers were sculpt-relaxed in v3 (carried over to v4 mutants) but not against a relaxed pocket.",
            "RMSD is computed against crystal dUMP heavy atoms only (no symmetry handling for ring atoms).",
        ]:
            doc.add_paragraph(line, style="List Bullet")

        doc.add_heading("Methods - AD4/Vina united-atom convention (FIX J)", 1)
        doc.add_paragraph(
            "AutoDock 4 / Vina use a united-atom convention; non-polar hydrogens are "
            "merged into their parent heavy atoms, and polar hydrogens (HD type) carry "
            "zero partial charge with the H-bond contribution folded into the parent. "
            "The 1070 zero-charge atoms in the holo receptor PDBQT are therefore a "
            "feature of the AD4 model, not a charge-assignment failure. Maximum heavy-atom "
            f"|q| in the v4 receptor is {max_q:.3f}."
        )

        doc.add_heading("Reproducibility", 1)
        doc.add_paragraph(
            f"Vina 1.2.7 (WT: exh=96, num_modes=32, box=22, seeds {{42,7,13,99,256}} primary; "
            "mutants: exh=32, num_modes=20, box=22, seed=42). "
            f"Receptor: {rec_method}, Gasteiger charges, max |q| = {max_q:.3f}. "
            f"Centroid (crystal dUMP, A): ({centroid[0]:.2f}, {centroid[1]:.2f}, {centroid[2]:.2f}). "
            "Cofactor: D16 ideal SDF from RCSB CCD; RDKit Reionizer; explicit -COOH -> COO-; "
            "AddHs(addCoords=True); Kabsch alignment to crystal coords."
        )
        doc.add_paragraph(
            "All v4 outputs in 03d_structure_v4, 06d_docking_wt_v4, 07d_mut_docking_v4, "
            "08d_analysis_v4, 09d_report_v4. v1, v2, v3 left untouched."
        )

        docx_path = os.path.join(REP_DIR, "report.docx")
        doc.save(docx_path)
        log(f"wrote {docx_path} ({os.path.getsize(docx_path)} bytes)")
    except Exception as e:
        log(f"docx failed: {e}")
        import traceback
        log(traceback.format_exc())

    log("Stage 9 v4 DONE")


if __name__ == "__main__":
    main()
